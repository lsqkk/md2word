"""Tests for ooxml_helpers.py — bookmarks, fields, SVG, numbering, borders, TOC."""

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from md2word.context import ConversionContext
from md2word.ooxml_helpers import (
    add_bookmark,
    add_bottom_border,
    add_num_pr,
    add_page_break,
    add_seq_number,
    add_toc,
    create_list_num_id,
    embed_svg_in_paragraph,
    ensure_list_abstract_nums,
    make_ref_field,
    set_outline_level,
    style_table,
)
from md2word.template import ParagraphFormat


class TestAddBookmark:
    def test_adds_bookmark_start_and_end(self):
        doc = Document()
        p = doc.add_paragraph("test")
        ctx = ConversionContext()
        name = add_bookmark(p, ctx, name="MyBookmark")
        assert name == "MyBookmark"

        xml = p._p.xml
        assert "w:bookmarkStart" in xml
        assert "w:bookmarkEnd" in xml

    def test_auto_generated_name(self):
        doc = Document()
        p = doc.add_paragraph("test")
        ctx = ConversionContext()
        name = add_bookmark(p, ctx)
        assert name.startswith("_Ref")

    def test_counter_increments(self):
        doc = Document()
        p = doc.add_paragraph("test")
        ctx = ConversionContext()
        add_bookmark(p, ctx)
        add_bookmark(p, ctx)
        assert ctx.bookmark_counter == 2


class TestMakeRefField:
    def test_returns_5_runs(self):
        runs = make_ref_field("参见上文", "MyHeading")
        assert len(runs) == 5
        # Check it contains instrText with REF
        xml_joined = "".join(r.xml for r in runs)
        assert "REF" in xml_joined
        assert "MyHeading" in xml_joined

    def test_different_display_text(self):
        runs = make_ref_field("Table 1", "table-1")
        xml_joined = "".join(r.xml for r in runs)
        assert "Table 1" in xml_joined
        assert "table-1" in xml_joined


class TestAddSeqNumber:
    def test_adds_seq_field(self):
        doc = Document()
        p = doc.add_paragraph()
        add_seq_number(p, "Equation")
        xml = p._p.xml
        assert "SEQ" in xml
        assert "Equation" in xml

    def test_default_seq_name(self):
        doc = Document()
        p = doc.add_paragraph()
        add_seq_number(p)
        xml = p._p.xml
        assert "Equation" in xml


class TestSetOutlineLevel:
    def test_sets_level(self):
        doc = Document()
        p = doc.add_paragraph()
        set_outline_level(p, 1)

        pPr = p._p.find(qn("w:pPr"))
        assert pPr is not None
        lvl = pPr.find(qn("w:outlineLvl"))
        assert lvl is not None
        assert lvl.get(qn("w:val")) == "0"

    def test_sets_heading_level_3(self):
        doc = Document()
        p = doc.add_paragraph()
        set_outline_level(p, 3)
        pPr = p._p.find(qn("w:pPr"))
        lvl = pPr.find(qn("w:outlineLvl"))
        assert lvl.get(qn("w:val")) == "2"


class TestAddBottomBorder:
    def test_adds_border_element(self):
        doc = Document()
        p = doc.add_paragraph()
        add_bottom_border(p, color="CC0000", sz="16")

        pPr = p._p.find(qn("w:pPr"))
        assert pPr is not None
        pBdr = pPr.find(qn("w:pBdr"))
        assert pBdr is not None
        bottom = pBdr.find(qn("w:bottom"))
        assert bottom is not None
        assert bottom.get(qn("w:color")) == "CC0000"
        assert bottom.get(qn("w:sz")) == "16"

    def test_default_values(self):
        doc = Document()
        p = doc.add_paragraph()
        add_bottom_border(p)

        pPr = p._p.find(qn("w:pPr"))
        pBdr = pPr.find(qn("w:pBdr"))
        bottom = pBdr.find(qn("w:bottom"))
        assert bottom.get(qn("w:color")) == "CCCCCC"
        assert bottom.get(qn("w:sz")) == "6"


class TestAddPageBreak:
    def test_adds_page_break(self):
        doc = Document()
        add_page_break(doc)
        # Should have at least one paragraph with <w:br w:type="page"/>
        found = False
        for p in doc.paragraphs:
            if "page" in p._p.xml:
                found = True
                break
        assert found


class TestAddToc:
    def test_adds_toc_field(self):
        doc = Document()
        styles = {
            "toc_title": ParagraphFormat(font_name="SimHei", bold=True,
                                          font_size_emu=18 * 12700),
        }
        add_toc(doc, styles, depth="1-3")
        xml = doc.element.xml
        assert "TOC" in xml
        assert '\\o "1-3"' in xml

    def test_minimal_styles(self):
        doc = Document()
        add_toc(doc, {}, depth="1-2")
        # Should not crash even without toc_title style
        xml = doc.element.xml
        assert "TOC" in xml


class TestStyleTable:
    def test_default_style_all_borders(self):
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        style_table(table, three_line=False)

        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        borders = tblPr.find(qn("w:tblBorders"))
        assert borders is not None
        # All 6 borders present
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = borders.find(qn(f"w:{edge}"))
            assert el is not None, f"Missing border: {edge}"

    def test_three_line_style(self):
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        style_table(table, three_line=True)

        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        borders = tblPr.find(qn("w:tblBorders"))
        # Check top, bottom, insideH have single style
        for edge in ("top", "bottom"):
            el = borders.find(qn(f"w:{edge}"))
            assert el is not None
            assert el.get(qn("w:val")) == "single"
        # Check left, right, insideV are none
        for edge in ("left", "right", "insideV"):
            el = borders.find(qn(f"w:{edge}"))
            assert el is not None
            assert el.get(qn("w:val")) == "none"


class TestEnsureListAbstractNums:
    def test_creates_abstract_nums(self):
        doc = Document()
        result = ensure_list_abstract_nums(doc)
        assert "bullet" in result
        assert "ordered" in result
        assert isinstance(result["bullet"], int)
        assert isinstance(result["ordered"], int)

    def test_cached(self):
        doc = Document()
        result1 = ensure_list_abstract_nums(doc)
        result2 = ensure_list_abstract_nums(doc)
        assert result1 == result2


class TestCreateListNumId:
    def test_creates_num_id(self):
        doc = Document()
        num_id = create_list_num_id(doc, "bullet")
        assert isinstance(num_id, int)
        assert num_id >= 1

    def test_ordered_type(self):
        doc = Document()
        num_id = create_list_num_id(doc, "ordered")
        assert isinstance(num_id, int)


class TestAddNumPr:
    def test_adds_numbering_properties(self):
        doc = Document()
        p = doc.add_paragraph()
        add_num_pr(p, num_id=1, ilvl=0)

        pPr = p._p.find(qn("w:pPr"))
        numPr = pPr.find(qn("w:numPr"))
        assert numPr is not None

        ilvl_el = numPr.find(qn("w:ilvl"))
        assert ilvl_el.get(qn("w:val")) == "0"
        num_id_el = numPr.find(qn("w:numId"))
        assert num_id_el.get(qn("w:val")) == "1"

    def test_replaces_existing(self):
        doc = Document()
        p = doc.add_paragraph()
        add_num_pr(p, num_id=1, ilvl=0)
        add_num_pr(p, num_id=2, ilvl=1)

        pPr = p._p.find(qn("w:pPr"))
        numPr_list = pPr.findall(qn("w:numPr"))
        assert len(numPr_list) == 1  # Replaced, not duplicated


class TestEmbedSvgInParagraph:
    def test_basic_svg_embed(self):
        """Minimal SVG embedding test."""
        doc = Document()
        p = doc.add_paragraph()
        svg_data = b'<svg xmlns="http://www.w3.org/2000/svg" width="100" height="50"></svg>'
        embed_svg_in_paragraph(doc, p, svg_data, alt="test svg", max_width_inches=4.0)
        xml = p._p.xml
        assert "svgBlip" in xml or "asvg" in xml or "SVG" in xml
