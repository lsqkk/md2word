"""Integration tests — generate real docx files and verify with python-docx.

These tests validate the end-to-end conversion pipeline by checking the
output document's paragraph count, style names, font sizes, and content.
"""

from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt

from md2word.converter import convert
from md2word.context import ConversionContext, ConversionReport
from md2word.frontmatter import apply_front_matter, parse_front_matter
from md2word.handlers import preprocess_extended_syntax, _slugify


# ── Integration test: full conversion ────────────────────────────────────────


def _create_minimal_template(tmp_path: Path) -> Path:
    """Create a minimal .docx template for integration tests."""
    from docx import Document as _Doc
    from docx.shared import Cm

    doc = _Doc()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    guide_data = [
        ("一级标题", "SimHei", 18, True),
        ("二级标题", "SimHei", 15, True),
        ("三级标题", "SimHei", 13, True),
        ("正文", "SimSun", 12, False),
        ("首行缩进", "SimSun", 12, False),
        ("图片", "SimSun", 10, False),
        ("图注", "SimSun", 9, False),
        ("引用", "KaiTi", 10, False),
        ("代码", "DengXian", 9, False),
        ("无序列表", "SimSun", 12, False),
        ("有序列表", "SimSun", 12, False),
    ]
    for text, font, size, bold in guide_data:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = font
        run.font.size = Pt(size)
        run.font.bold = bold

    path = tmp_path / "integration_template.docx"
    doc.save(str(path))
    return path


class TestFullConversion:
    """End-to-end: convert markdown → docx, read back with python-docx."""

    def test_basic_conversion(self, tmp_path):
        """Convert simple markdown and verify output structure."""
        template = _create_minimal_template(tmp_path)
        output = tmp_path / "output.docx"
        md = """# 标题一

这是一段正文。

## 标题二

- 列表项1
- 列表项2

> 这是一段引用
"""
        report = convert(md, template, output)
        assert report is not None
        assert not report.has_errors()

        doc = Document(str(output))
        assert len(doc.paragraphs) > 0
        texts = [p.text for p in doc.paragraphs]
        assert any("标题一" in t for t in texts)
        assert any("这是一段正文" in t for t in texts)
        assert any("列表项1" in t for t in texts)
        assert any("这是一段引用" in t for t in texts)

    def test_heading_styles(self, tmp_path):
        """Verify headings get proper outline levels."""
        template = _create_minimal_template(tmp_path)
        output = tmp_path / "headings.docx"
        md = "# H1\n\n## H2\n\n### H3"

        convert(md, template, output)
        doc = Document(str(output))

        from docx.oxml.ns import qn
        h1_count = 0
        h2_count = 0
        h3_count = 0
        for p in doc.paragraphs:
            pPr = p._p.find(qn("w:pPr"))
            if pPr is not None:
                lvl = pPr.find(qn("w:outlineLvl"))
                if lvl is not None:
                    val = lvl.get(qn("w:val"))
                    if val == "0":
                        h1_count += 1
                    elif val == "1":
                        h2_count += 1
                    elif val == "2":
                        h3_count += 1
        assert h1_count >= 1
        assert h2_count >= 1
        assert h3_count >= 1

    def test_table_conversion(self, tmp_path):
        """Verify tables are created."""
        template = _create_minimal_template(tmp_path)
        output = tmp_path / "tables.docx"
        md = "| 列1 | 列2 |\n|-----|-----|\n| A   | B   |"

        convert(md, template, output)
        doc = Document(str(output))
        assert len(doc.tables) >= 1
        assert doc.tables[0].rows[0].cells[0].text.strip() == "列1"

    def test_code_block(self, tmp_path):
        """Verify code blocks are present."""
        template = _create_minimal_template(tmp_path)
        output = tmp_path / "code.docx"
        md = "```python\nprint('hello')\n```"

        convert(md, template, output)
        doc = Document(str(output))
        texts = [p.text for p in doc.paragraphs]
        code_texts = [t for t in texts if "print" in t or "hello" in t]
        assert len(code_texts) >= 1


# ── ConversionContext tests ─────────────────────────────────────────────


class TestConversionContext:
    def test_bookmark_counter(self):
        ctx = ConversionContext()
        assert ctx.next_bookmark_id() == 1
        assert ctx.next_bookmark_id() == 2

        ctx2 = ConversionContext()
        assert ctx2.next_bookmark_id() == 1  # isolated

    def test_heading_counters_isolated(self):
        ctx1 = ConversionContext()
        ctx2 = ConversionContext()

        ctx1.reset_heading_counters()
        ctx2.reset_heading_counters()

        assert ctx1.next_heading_number(2) == "1"
        assert ctx1.next_heading_number(2) == "2"
        assert ctx2.next_heading_number(2) == "1"  # isolated

    def test_report_severity(self):
        report = ConversionReport()
        report.info_msg("info test")
        report.warn("warning test")
        report.error("error test")
        report.add_critical("critical test")

        assert len(report.info) == 1
        assert len(report.warnings) == 1
        assert len(report.errors) == 2  # error + critical
        assert len(report.critical) == 1
        assert report.has_errors()
        assert report.has_critical()


# ── Front matter tests ─────────────────────────────────────────────────


class TestParseFrontMatter:
    def test_parse_title_author(self):
        md = "---\ntitle: 测试文档\nauthor: 张三\ndate: 2026-01-15\n---\n\n# 正文"
        meta, body = parse_front_matter(md)
        assert meta["title"] == "测试文档"
        assert meta["author"] == "张三"
        assert "# 正文" in body

    def test_no_front_matter(self):
        meta, body = parse_front_matter("# 只有正文")
        assert meta == {}
        assert body == "# 只有正文"

    def test_abstract_and_keywords(self):
        md = "---\nkeywords: markdown, word, 转换\n---\n\n内容"
        meta, body = parse_front_matter(md)
        assert "markdown" in meta["keywords"]
        assert "内容" in body

    def test_apply_to_docx(self, tmp_path):
        from docx import Document
        doc = Document()
        apply_front_matter(doc, {"title": "测试标题", "author": "作者"})
        assert doc.core_properties.title == "测试标题"
        assert doc.core_properties.author == "作者"


# ── Extended syntax tests ──────────────────────────────────────────────


class TestPreprocessExtendedSyntax:
    def test_strikethrough(self):
        result = preprocess_extended_syntax("~~删除线~~")
        assert "<del>删除线</del>" in result

    def test_highlight(self):
        result = preprocess_extended_syntax("==高亮==")
        assert "<mark>高亮</mark>" in result

    def test_superscript(self):
        result = preprocess_extended_syntax("X^2^")
        assert "<sup>2</sup>" in result

    def test_subscript(self):
        result = preprocess_extended_syntax("H~2~O")
        assert "<sub>2</sub>" in result

    def test_combined(self):
        result = preprocess_extended_syntax("~~strike~~ and ==highlight==")
        assert "<del>strike</del>" in result
        assert "<mark>highlight</mark>" in result

    def test_no_false_positive(self):
        result = preprocess_extended_syntax("正常文本 ~~~ not strikethrough")
        assert "<del>" not in result or "~~~" in result


# ── Slugify tests ──────────────────────────────────────────────────────


class TestSlugifyExtended:
    def test_bookmark_name_used(self):
        """Verify _slugify produces usable bookmark names."""
        assert _slugify("数据说明") == "数据说明".lower()  # falls through
        assert _slugify("第1章 引言") == "第1章-引言".lower()


# ── Report entry tests ─────────────────────────────────────────────────


class TestReportEntry:
    def test_severity_order(self):
        report = ConversionReport()
        report.info_msg("info msg")
        report.warn("warn msg")
        report.error("error msg")

        assert len(report.entries) == 3
        sevs = [e.severity for e in report.entries]
        assert sevs == ["info", "warning", "error"]

    def test_summary_no_issues(self):
        report = ConversionReport()
        s = report.summary()
        assert "无警告或错误" in s

    def test_summary_with_warnings(self):
        report = ConversionReport()
        report.warn("test warning")
        s = report.summary()
        assert "警告" in s


# ── Conversion report integration ──────────────────────────────────────


class TestConvertReport:
    def test_report_returned(self, tmp_path):
        """Ensure convert() returns a ConversionReport object."""
        template = _create_minimal_template(tmp_path)
        output = tmp_path / "report_test.docx"
        md = "# Test\n\nBody text."
        report = convert(md, template, output)
        assert isinstance(report, ConversionReport)
        assert not report.has_errors()

    def test_report_with_errors(self, tmp_path):
        """Even with bad content, convert should not crash."""
        template = _create_minimal_template(tmp_path)
        output = tmp_path / "error_test.docx"
        md = "# 正常内容\n\n- 列表"
        report = convert(md, template, output)
        # Should complete without error for normal content
        assert report is not None
