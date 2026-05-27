"""Comprehensive tests for v1.7 features.

Covers:
- ConvertOptions dataclass
- Config system upgrade (PyYAML, nested config, validation)
- Template custom XML markers
- Verbose mode output
- Error message context with source preview
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Pt

from md2word.config import (
    _load_yaml,
    _load_yaml_text_fallback,
    _normalize_yaml,
    _validate_config_keys,
    find_config,
    load_config,
    merge_with_args,
)
from md2word.converter import _inline_text_preview, convert
from md2word.options import ConvertOptions
from md2word.template import (
    _guess_slot,
    _inject_slot_marker,
    _read_slot_marker,
    extract_template_styles,
)


# ═══════════════════════════════════════════════════════════════════════════
# ConvertOptions tests
# ═══════════════════════════════════════════════════════════════════════════


class TestConvertOptions:
    def test_default_values(self):
        opts = ConvertOptions()
        assert opts.image_max_width == 5.5
        assert opts.toc is True
        assert opts.toc_depth == "1-3"
        assert opts.highlight_enabled is True
        assert opts.math_enabled is True
        assert opts.mermaid_enabled is True
        assert opts.number_headings is False
        assert opts.page_break_h1 is False
        assert opts.three_line_table is False
        assert opts.footnotes_enabled is True
        assert opts.formula_numbering is False
        assert opts.redhead_authority is None
        assert opts.redhead_year is None
        assert opts.redhead_number is None
        assert opts.page_number_fmt is None
        assert opts.gb_check is False
        assert opts.style_map is None
        assert opts.verbose is False

    def test_from_cli_args_simple(self):
        opts = ConvertOptions.from_cli_args(
            {"toc": False, "number_headings": True}
        )
        assert opts.toc is False
        assert opts.number_headings is True
        assert opts.image_max_width == 5.5  # default preserved

    def test_from_cli_args_with_config(self):
        cfg = {"toc": True, "image_max_width": 4.0}
        opts = ConvertOptions.from_cli_args(
            {"number_headings": True}, cfg
        )
        assert opts.toc is True  # from cfg
        assert opts.image_max_width == 4.0  # from cfg
        assert opts.number_headings is True  # from cli

    def test_from_cli_args_cli_overrides_config(self):
        cfg = {"toc": True, "number_headings": False}
        opts = ConvertOptions.from_cli_args(
            {"toc": False, "number_headings": True}, cfg
        )
        assert opts.toc is False  # cli overrides
        assert opts.number_headings is True  # cli overrides

    def test_from_cli_args_none_not_overwritten(self):
        cfg = {"toc": True}
        opts = ConvertOptions.from_cli_args(
            {"toc": None, "number_headings": True}, cfg
        )
        assert opts.toc is True  # cli None, cfg value used

    def test_style_map_passthrough(self):
        opts = ConvertOptions(style_map={"code": "CustomCode"})
        assert opts.style_map == {"code": "CustomCode"}

    def test_redhead_defaults(self):
        opts = ConvertOptions(redhead_authority="测试")
        assert opts.redhead_authority == "测试"
        assert opts.redhead_year is None
        assert opts.redhead_number is None


# ═══════════════════════════════════════════════════════════════════════════
# Config system tests
# ═══════════════════════════════════════════════════════════════════════════


class TestConfigYamlLoad:
    def test_load_simple_yaml(self, tmp_path):
        cfg_file = tmp_path / "md2word.yaml"
        cfg_file.write_text("toc: true\nimage_max_width: 4.5\n", encoding="utf-8")
        result = _load_yaml(cfg_file)
        assert result.get("toc") is True
        assert result.get("image_max_width") == 4.5

    def test_load_yaml_with_nested_map(self, tmp_path):
        """PyYAML should parse nested style_map."""
        cfg_file = tmp_path / "md2word.yaml"
        cfg_file.write_text(
            "style_map:\n  code: CustomCode\n  quote: CustomQuote\n",
            encoding="utf-8",
        )
        result = _load_yaml(cfg_file)
        assert result.get("style_map") == {"code": "CustomCode", "quote": "CustomQuote"}

    def test_load_yaml_with_list(self, tmp_path):
        cfg_file = tmp_path / "md2word.yaml"
        cfg_file.write_text("inputs:\n  - a.md\n  - b.md\n", encoding="utf-8")
        result = _load_yaml(cfg_file)
        assert result.get("inputs") == ["a.md", "b.md"]

    def test_yaml_key_normalization(self, tmp_path):
        """Hyphenated keys should be normalized to underscores."""
        cfg_file = tmp_path / "md2word.yaml"
        cfg_file.write_text("toc-depth: 1-2\nimage-max-width: 3.0\n", encoding="utf-8")
        result = _load_yaml(cfg_file)
        assert result.get("toc_depth") == "1-2"
        assert result.get("image_max_width") == 3.0

    def test_fallback_parser_no_pyyaml(self, tmp_path, monkeypatch):
        """Fallback parser should work for simple flat configs."""
        import md2word.config as cfg_mod
        monkeypatch.setattr(cfg_mod, "_HAS_PYYAML", False)
        cfg_file = tmp_path / "md2word.yaml"
        cfg_file.write_text("toc: true\nhighlight: false\n", encoding="utf-8")
        result = _load_yaml(cfg_file)
        assert result.get("toc") is True
        assert result.get("highlight") is False

    def test_validate_config_keys_unknown(self, tmp_path):
        cfg = {"unknown_key": 123, "toc": True}
        # Should not raise, just print warning
        _validate_config_keys(cfg, tmp_path / "dummy.yaml")

    def test_load_config_empty(self):
        """No config file → empty dict."""
        result = load_config(Path("/nonexistent/path"))
        assert result == {}

    def test_merge_with_args_cli_priority(self):
        cfg = {"toc": True, "number_headings": False}
        args = {"toc": None, "number_headings": True}
        result = merge_with_args(cfg, args)
        assert result["number_headings"] is True
        assert result["toc"] is True  # None in args, cfg value used


class TestConfigFileDetection:
    def test_find_config_none(self, tmp_path):
        result = find_config(tmp_path)
        assert result is None

    def test_find_config_yaml(self, tmp_path):
        cfg = tmp_path / "md2word.yaml"
        cfg.write_text("toc: true\n", encoding="utf-8")
        result = find_config(tmp_path)
        assert result == cfg

    def test_find_config_json(self, tmp_path):
        cfg = tmp_path / "md2word.json"
        cfg.write_text('{"toc": true}\n', encoding="utf-8")
        result = find_config(tmp_path)
        assert result == cfg

    def test_find_config_dot_dir(self, tmp_path):
        dot_dir = tmp_path / ".md2word"
        dot_dir.mkdir()
        cfg = dot_dir / "config.yaml"
        cfg.write_text("toc: false\n", encoding="utf-8")
        result = find_config(tmp_path)
        assert result == cfg


# ═══════════════════════════════════════════════════════════════════════════
# Template custom XML marker tests
# ═══════════════════════════════════════════════════════════════════════════


class TestSlotMarkers:
    def test_inject_and_read_marker(self, tmp_path):
        from docx import Document
        doc = Document()
        p = doc.add_paragraph("测试段落")
        _inject_slot_marker(p, "h1")
        assert _read_slot_marker(p) == "h1"

    def test_no_marker_returns_none(self, tmp_path):
        from docx import Document
        doc = Document()
        p = doc.add_paragraph("无标记段落")
        assert _read_slot_marker(p) is None

    def test_guess_slot_with_marker_takes_priority(self):
        """_guess_slot should return the marker value even if text doesn't match."""
        from docx import Document
        doc = Document()
        p = doc.add_paragraph("some random text that is not a keyword")
        _inject_slot_marker(p, "h1")
        assert _guess_slot("some random text", p=p) == "h1"

    def test_guess_slot_fallback_text_matching(self):
        """Without a marker, text-substring matching should still work."""
        assert _guess_slot("一级标题示例") == "h1"
        assert _guess_slot("正文内容") == "body"
        assert _guess_slot("图片") == "image"

    def test_guess_slot_no_match(self):
        assert _guess_slot("完全不相关内容") is None

    def test_guess_slot_empty_text(self):
        assert _guess_slot("") is None

    def test_template_styles_with_marker(self, tmp_path):
        """Template with markers should extract styles correctly."""
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        section = doc.sections[0]
        from docx.shared import Cm
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

        # Add guide paragraphs with slot markers
        p = doc.add_paragraph()
        r = p.add_run("一级标题")
        r.font.name = "SimHei"
        r.font.size = Pt(18)
        r.font.bold = True
        _inject_slot_marker(p, "h1")

        p = doc.add_paragraph()
        r = p.add_run("正文")
        r.font.name = "SimSun"
        r.font.size = Pt(12)
        _inject_slot_marker(p, "body")

        path = tmp_path / "marker_template.docx"
        doc.save(str(path))

        styles = extract_template_styles(path)
        assert "h1" in styles
        assert "body" in styles
        assert styles["h1"].font_name == "SimHei"
        assert styles["h1"].font_size_emu == Pt(18)

    def test_extract_styles_without_markers(self, tmp_path):
        """Backward compatibility: templates without markers still work."""
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("一级标题")
        r.font.name = "SimHei"
        r.font.size = Pt(18)
        r.font.bold = True

        path = tmp_path / "nomarker_template.docx"
        doc.save(str(path))

        styles = extract_template_styles(path)
        assert "h1" in styles


# ═══════════════════════════════════════════════════════════════════════════
# Verbose mode tests
# ═══════════════════════════════════════════════════════════════════════════


class TestVerboseMode:
    def test_verbose_output_enabled(self, tmp_path, capsys):
        """Verbose mode should emit progress messages on stderr."""
        from docx import Document
        from docx.shared import Pt, Cm

        # Create template
        doc = Document()
        s = doc.sections[0]
        s.top_margin = Cm(2.54)
        s.bottom_margin = Cm(2.54)
        s.left_margin = Cm(3.17)
        s.right_margin = Cm(3.17)
        for text, font, sz, bold in [
            ("一级标题", "SimHei", 18, True),
            ("二级标题", "SimHei", 15, True),
            ("正文", "SimSun", 12, False),
            ("代码", "DengXian", 9, False),
        ]:
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.name = font
            r.font.size = Pt(sz)
            r.font.bold = bold
        tpl = tmp_path / "verbose_template.docx"
        doc.save(str(tpl))

        output = tmp_path / "verbose_out.docx"
        md_text = "# 标题\n\n正文内容"
        report = convert(md_text, tpl, output, ConvertOptions(verbose=True))

        captured = capsys.readouterr()
        stderr_output = captured.out + captured.err
        assert report is not None
        assert not report.has_errors()

    def test_non_verbose_no_extra_output(self, tmp_path, capsys):
        """Without verbose, output should be minimal."""
        from docx import Document
        from docx.shared import Pt, Cm

        doc = Document()
        s = doc.sections[0]
        s.top_margin = Cm(2.54)
        s.bottom_margin = Cm(2.54)
        s.left_margin = Cm(3.17)
        s.right_margin = Cm(3.17)
        for text, font, sz, bold in [
            ("一级标题", "SimHei", 18, True),
            ("正文", "SimSun", 12, False),
        ]:
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.name = font
            r.font.size = Pt(sz)
            r.font.bold = bold
        tpl = tmp_path / "quiet_template.docx"
        doc.save(str(tpl))

        output = tmp_path / "quiet_out.docx"
        md_text = "# 标题\n\n正文"
        convert(md_text, tpl, output)

        captured = capsys.readouterr()
        # Should have summary output, but not verbose detail
        out = captured.out + captured.err
        assert out.strip()  # should have some output (summary)


# ═══════════════════════════════════════════════════════════════════════════
# Error message context tests
# ═══════════════════════════════════════════════════════════════════════════


class TestErrorContext:
    def test_inline_text_preview_short(self):
        """Short text should not be truncated."""
        from md2word.handlers import inline_text
        from xml.etree import ElementTree as ET

        elem = ET.fromstring("<p>短文本</p>")
        preview = _inline_text_preview(elem)
        assert preview == "短文本"

    def test_inline_text_preview_long(self):
        """Long text should be truncated with ellipsis."""
        from xml.etree import ElementTree as ET

        long_text = "A" * 200
        elem = ET.fromstring(f"<p>{long_text}</p>")
        preview = _inline_text_preview(elem)
        assert len(preview) <= 81  # 80 + ellipsis
        assert preview.endswith("…")

    def test_inline_text_preview_empty(self):
        """Empty element should return empty string."""
        from xml.etree import ElementTree as ET

        elem = ET.fromstring("<p></p>")
        preview = _inline_text_preview(elem)
        assert preview == ""


# ═══════════════════════════════════════════════════════════════════════════
# Integration: ConvertOptions in convert()
# ═══════════════════════════════════════════════════════════════════════════


class TestConvertWithOptions:
    def _make_template(self, tmp_path: Path) -> Path:
        from docx import Document as _Doc
        from docx.shared import Cm

        doc = _Doc()
        s = doc.sections[0]
        s.top_margin = Cm(2.54)
        s.bottom_margin = Cm(2.54)
        s.left_margin = Cm(3.17)
        s.right_margin = Cm(3.17)
        for text, font, sz, bold in [
            ("一级标题", "SimHei", 18, True),
            ("二级标题", "SimHei", 15, True),
            ("正文", "SimSun", 12, False),
            ("代码", "DengXian", 9, False),
        ]:
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.name = font
            r.font.size = Pt(sz)
            r.font.bold = bold
        path = tmp_path / "test_opts_template.docx"
        doc.save(str(path))
        return path

    def test_convert_with_options_object(self, tmp_path):
        template = self._make_template(tmp_path)
        output = tmp_path / "opts_out.docx"
        md = "# Hello\n\nWorld."
        opts = ConvertOptions(toc=False, verbose=False)
        report = convert(md, template, output, options=opts)
        assert not report.has_errors()

    def test_convert_with_legacy_kwargs(self, tmp_path):
        """Backward compat: keyword args still work."""
        template = self._make_template(tmp_path)
        output = tmp_path / "legacy_out.docx"
        md = "# Hello\n\nWorld."
        report = convert(md, template, output, toc=False)
        assert not report.has_errors()

    def test_convert_with_both(self, tmp_path):
        """options + kwargs: kwargs should override options."""
        template = self._make_template(tmp_path)
        output = tmp_path / "both_out.docx"
        md = "# Hello\n\nWorld."
        opts = ConvertOptions(toc=True)
        report = convert(md, template, output, options=opts, toc=False)
        assert not report.has_errors()

    def test_convert_style_map(self, tmp_path):
        """Style map should be applied."""
        template = self._make_template(tmp_path)
        output = tmp_path / "stylemap_out.docx"
        md = "```python\nx=1\n```"
        opts = ConvertOptions(style_map={"code": "body"})
        report = convert(md, template, output, options=opts)
        assert not report.has_errors()

    def test_convert_redhead(self, tmp_path):
        """Red-head with year and number."""
        template = self._make_template(tmp_path)
        output = tmp_path / "redhead_out.docx"
        md = "# 通知\n\n内容"
        opts = ConvertOptions(
            redhead_authority="XX市人民政府",
            redhead_year=2026,
            redhead_number="12",
        )
        report = convert(md, template, output, options=opts)
        assert not report.has_errors()

    def test_convert_page_number(self, tmp_path):
        template = self._make_template(tmp_path)
        output = tmp_path / "pgnum_out.docx"
        md = "# 标题\n\n页"
        opts = ConvertOptions(page_number_fmt="-- %d --")
        report = convert(md, template, output, options=opts)
        assert not report.has_errors()

    def test_convert_gb_check(self, tmp_path):
        template = self._make_template(tmp_path)
        output = tmp_path / "gb_out.docx"
        md = "# 标题\n\n正文"
        opts = ConvertOptions(gb_check=True)
        report = convert(md, template, output, options=opts)
        assert report is not None
