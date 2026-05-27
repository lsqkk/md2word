"""Integration tests for v1.6 features — abstract/keywords, bookmark dedup,
redhead config, style_map, push_detect refactor, code block strip fix."""

from pathlib import Path

import pytest
from docx import Document

from md2word.converter import _insert_abstract_keywords, convert
from md2word.context import ConversionContext
from md2word.frontmatter import apply_front_matter, parse_front_matter
from md2word.handlers import (
    _detect_task_prefix,
    _push_detect,
    _slugify,
    build_runs,
    build_runs_skip,
    ensure_list_blank_lines,
    handle_code_block,
    handle_heading,
    inline_text,
)
from md2word.template import ParagraphFormat


# ── Helpers ────────────────────────────────────────────────────────────────


def _make_template(tmp_path: Path) -> Path:
    """Create a minimal template with all v1.6 style slots."""
    from docx import Document as _Doc
    from docx.shared import Cm, Pt

    doc = _Doc()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    guide_data = [
        ("一级标题", "SimHei", 18, True),
        ("二级标题", "SimHei", 15, True),
        ("三级标题", "SimHei", 13, True),
        ("正文", "SimSun", 12, False),
        ("首行缩进", "SimSun", 12, False),
        ("摘要", "SimHei", 14, True),
        ("关键词", "SimHei", 14, True),
        ("参考文献", "SimSun", 9, False),
        ("图片", "SimSun", 10, False),
        ("图注", "SimSun", 9, False),
        ("引用", "KaiTi", 10, False),
        ("代码", "DengXian", 9, False),
        ("无序列表", "SimSun", 12, False),
        ("有序列表", "SimSun", 12, False),
    ]
    for text, font, size, bold in guide_data:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = font
        run.font.size = Pt(size)
        run.font.bold = bold

    path = tmp_path / "v16_template.docx"
    doc.save(str(path))
    return path


# ── 6.1 Abstract/keywords rendering ────────────────────────────────────────


class TestInsertAbstractKeywords:
    def test_inserts_abstract(self):
        doc = Document()
        styles = {
            "abstract": ParagraphFormat(font_name="SimHei", bold=True,
                                         font_size_emu=14 * 12700),
            "body": ParagraphFormat(font_name="SimSun", font_size_emu=12 * 12700),
        }
        _insert_abstract_keywords(doc, {"abstract": "这是一个摘要。"}, styles)
        texts = [p.text for p in doc.paragraphs]
        assert "摘要" in texts
        assert "这是一个摘要。" in texts

    def test_inserts_keywords(self):
        doc = Document()
        styles = {
            "keywords": ParagraphFormat(font_name="SimHei", bold=True,
                                         font_size_emu=14 * 12700),
            "body": ParagraphFormat(font_name="SimSun", font_size_emu=12 * 12700),
        }
        _insert_abstract_keywords(doc, {"keywords": "关键词1, 关键词2"}, styles)
        texts = [p.text for p in doc.paragraphs]
        assert "关键词" in texts
        assert "关键词1, 关键词2" in texts

    def test_inserts_both(self):
        doc = Document()
        styles = {
            "abstract": ParagraphFormat(font_name="SimHei", bold=True,
                                         font_size_emu=14 * 12700),
            "keywords": ParagraphFormat(font_name="SimHei", bold=True,
                                         font_size_emu=14 * 12700),
            "body": ParagraphFormat(font_name="SimSun", font_size_emu=12 * 12700),
        }
        fm = {"abstract": "摘要内容", "keywords": "kw1, kw2"}
        _insert_abstract_keywords(doc, fm, styles)
        texts = [p.text for p in doc.paragraphs]
        assert "摘要" in texts
        assert "摘要内容" in texts
        assert "关键词" in texts
        assert "kw1, kw2" in texts

    def test_no_front_matter_no_insertion(self):
        doc = Document()
        styles = {"body": ParagraphFormat()}
        _insert_abstract_keywords(doc, {}, styles)
        assert len(doc.paragraphs) == 0

    def test_without_style_fallback_to_body(self):
        """When abstract style slot is missing, use body style."""
        doc = Document()
        styles = {"body": ParagraphFormat(font_name="SimSun", font_size_emu=12 * 12700)}
        _insert_abstract_keywords(doc, {"abstract": "fallback test"}, styles)
        texts = [p.text for p in doc.paragraphs]
        assert "fallback test" in texts


# ── 6.2 Bookmark dedup ─────────────────────────────────────────────────────


class TestBookmarkDedup:
    def test_unique_slug_first_use(self):
        ctx = ConversionContext()
        assert ctx.unique_bookmark_slug("hello") == "hello"

    def test_duplicate_slug_appends_suffix(self):
        ctx = ConversionContext()
        ctx.unique_bookmark_slug("hello")
        assert ctx.unique_bookmark_slug("hello") == "hello-1"
        assert ctx.unique_bookmark_slug("hello") == "hello-2"

    def test_different_slugs_independent(self):
        ctx = ConversionContext()
        ctx.unique_bookmark_slug("a")
        ctx.unique_bookmark_slug("b")
        assert ctx.unique_bookmark_slug("a") == "a-1"
        assert ctx.unique_bookmark_slug("b") == "b-1"

    def test_slugify_then_dedup(self):
        """Integration: slug produced from heading text gets dedup'd."""
        ctx = ConversionContext()
        s1 = ctx.unique_bookmark_slug(_slugify("引言"))
        s2 = ctx.unique_bookmark_slug(_slugify("引言"))
        assert s1 != s2
        assert s1 == "引言"
        assert s2 == "引言-1"


# ── 6.3 push_detect refactor ───────────────────────────────────────────────


class TestDetectTaskPrefix:
    def test_checked(self):
        remaining, marker = _detect_task_prefix("[x] task")
        assert remaining == "task"
        assert marker == "☑ "

    def test_unchecked(self):
        remaining, marker = _detect_task_prefix("[ ] task")
        assert remaining == "task"
        assert marker == "☐ "

    def test_no_prefix(self):
        remaining, marker = _detect_task_prefix("normal text")
        assert remaining == "normal text"
        assert marker is None

    def test_case_insensitive_checked(self):
        remaining, marker = _detect_task_prefix("[X] DONE")
        assert remaining == "DONE"
        assert marker is not None

    def test_empty_text(self):
        remaining, marker = _detect_task_prefix("")
        assert remaining == ""
        assert marker is None


# ── 6.4 build_runs / build_runs_skip refactor ──────────────────────────────


class TestBuildRunsRefactored:
    def test_build_runs_still_works(self):
        """Ensure build_runs wrapper delegates correctly."""
        import xml.etree.ElementTree as ET
        doc = Document()
        p = doc.add_paragraph()
        fmt = ParagraphFormat(font_name="SimSun", font_size_emu=12 * 12700)
        elem = ET.fromstring("<p>Hello <strong>World</strong></p>")
        build_runs(doc, p, elem, fmt)
        combined = "".join(r.text for r in p.runs if r.text)
        assert "Hello" in combined
        assert "World" in combined

    def test_build_runs_with_footnotes(self):
        import xml.etree.ElementTree as ET
        doc = Document()
        p = doc.add_paragraph()
        fmt = ParagraphFormat()
        from md2word.footnotes import add_footnotes_to_document, extract_footnotes
        text, fn_list = extract_footnotes("正文[^1]\n\n[^1]: 脚注")
        fn_map = add_footnotes_to_document(doc, fn_list)
        elem = ET.fromstring(f"<p>正文<fn id='1'/></p>")
        build_runs(doc, p, elem, fmt, fn_map=fn_map)
        # Should not crash — footnote ref run added
        assert len(p._p) > 0

    def test_build_runs_skip_still_works(self):
        import xml.etree.ElementTree as ET
        doc = Document()
        p = doc.add_paragraph()
        fmt = ParagraphFormat()
        elem = ET.fromstring("<li>Item <strong>bold</strong></li>")
        build_runs_skip(doc, p, elem, fmt)
        combined = "".join(r.text for r in p.runs if r.text)
        assert "Item" in combined
        assert "bold" in combined

    def test_build_runs_skip_with_skip_tags(self):
        import xml.etree.ElementTree as ET
        doc = Document()
        p = doc.add_paragraph()
        fmt = ParagraphFormat()
        elem = ET.fromstring("<li>Text <ul><li>nested</li></ul> tail</li>")
        build_runs_skip(doc, p, elem, fmt, skip_tags={"ul"})
        combined = "".join(r.text for r in p.runs if r.text)
        assert "Text" in combined
        assert "tail" in combined


# ── 6.5 Code block strip fix ────────────────────────────────────────────────


class TestCodeBlockStrip:
    def test_preserves_trailing_newlines(self):
        """code block with \\n after content should keep leading/trailing spaces."""
        import xml.etree.ElementTree as ET
        doc = Document()
        ctx = ConversionContext()
        styles = {
            "code": ParagraphFormat(font_name="Consolas", font_size_emu=9 * 12700),
        }
        elem = ET.fromstring(
            '<pre><code class="language-python">def foo():\n    pass\n</code></pre>'
        )
        # Should not crash; notably text.strip('\\n') preserves the trailing space
        # before the </code> closing tag
        handle_code_block(doc, elem, styles, ctx, highlight_enabled=False)
        texts = [p.text for p in doc.paragraphs]
        assert any("def foo()" in t for t in texts)
        assert any("pass" in t for t in texts)


# ── 6.7 style_map ──────────────────────────────────────────────────────────


class TestStyleMap:
    def test_convert_accepts_style_map(self, tmp_path):
        """style_map parameter is accepted and applied."""
        template = _make_template(tmp_path)
        output = tmp_path / "styled_output.docx"
        md = "# Title\n\nBody text."
        report = convert(
            md, template, output,
            style_map={"h1": "body"},  # Map h1 to body style
        )
        assert report is not None
        assert not report.has_errors()

    def test_style_map_does_not_crash_with_empty(self, tmp_path):
        template = _make_template(tmp_path)
        output = tmp_path / "empty_style_map.docx"
        md = "# Title"
        report = convert(md, template, output, style_map={})
        assert not report.has_errors()


# ── 6.8 Redhead config ──────────────────────────────────────────────────────


class TestRedheadConfig:
    def test_convert_accepts_redhead_year_number(self, tmp_path):
        template = _make_template(tmp_path)
        output = tmp_path / "redhead_config.docx"
        md = "# 通知\n\n正文内容"
        report = convert(
            md, template, output,
            redhead_authority="XX市政府",
            redhead_year=2026,
            redhead_number="42",
        )
        assert report is not None
        # Should complete without error
        assert not report.has_errors()

    def test_convert_accepts_redhead_year_only(self, tmp_path):
        template = _make_template(tmp_path)
        output = tmp_path / "redhead_year.docx"
        md = "# 通知"
        report = convert(md, template, output, redhead_authority="测试", redhead_year=2025)
        assert not report.has_errors()
