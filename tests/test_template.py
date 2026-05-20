"""Tests for template style extraction."""

import pytest
from docx import Document
from docx.shared import Pt

from md2word.template import (
    ParagraphFormat,
    extract_template_styles,
    validate_template,
    _guess_slot,
    _STYLE_KEYWORDS,
)


class TestGuessSlot:
    def test_known_keywords(self):
        assert _guess_slot("一级标题") == "h1"
        assert _guess_slot("二级标题") == "h2"
        assert _guess_slot("三级标题") == "h3"
        assert _guess_slot("正文") == "body"
        assert _guess_slot("首行缩进") == "body_indent"
        assert _guess_slot("图片") == "image"
        assert _guess_slot("图注") == "figcaption"
        assert _guess_slot("引用") == "quote"
        assert _guess_slot("代码") == "code"
        assert _guess_slot("无序列表") == "bullet_list"
        assert _guess_slot("有序列表") == "number_list"
        assert _guess_slot("目录标题") == "toc_title"

    def test_english_keywords(self):
        assert _guess_slot("heading 1") == "h1"
        assert _guess_slot("ordered list") == "number_list"
        assert _guess_slot("table of contents") == "toc_title"

    def test_ignores_punctuation_and_whitespace(self):
        assert _guess_slot("  一级标题  ") == "h1"
        assert _guess_slot("【正文】") == "body"

    def test_unknown_text_returns_none(self):
        assert _guess_slot("完全无关的文字") is None
        assert _guess_slot("") is None


class TestParagraphFormat:
    def test_default_values(self):
        pf = ParagraphFormat()
        assert pf.font_name is None
        assert pf.bold is None
        assert pf.font_size_emu is None

    def test_from_docx_paragraph(self):
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("测试")
        run.font.name = "SimSun"
        run.font.size = Pt(12)
        run.font.bold = True

        pf = ParagraphFormat.from_docx_paragraph(p)
        assert pf.font_name == "SimSun"
        assert pf.font_size_emu == 12 * 12700  # pt → EMU
        assert pf.bold is True

    def test_from_empty_paragraph(self):
        doc = Document()
        p = doc.add_paragraph()
        pf = ParagraphFormat.from_docx_paragraph(p)
        assert pf.font_name is None

    def test_apply_to_run(self):
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("test")

        pf = ParagraphFormat(font_name="SimHei", font_size_emu=14 * 12700, bold=True)
        pf.apply_to_run(run)

        assert run.font.name == "SimHei"
        assert run.font.bold is True
        assert run.font.size == Pt(14)

    def test_roundtrip_preserves_format(self):
        """Apply formatting then extract it back."""
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("test")
        run.font.name = "KaiTi"
        run.font.size = Pt(10.5)
        run.font.bold = False

        pf1 = ParagraphFormat.from_docx_paragraph(p)

        p2 = doc.add_paragraph()
        r2 = p2.add_run("result")
        pf1.apply_to_run(r2)

        pf2 = ParagraphFormat.from_docx_paragraph(p2)
        assert pf2.font_name == pf1.font_name
        assert pf2.bold == pf1.bold


class TestExtractStyles:
    def test_extracts_all_guide_styles(self, minimal_template):
        styles = extract_template_styles(minimal_template)
        assert "h1" in styles
        assert "h2" in styles
        assert "h3" in styles
        assert "body" in styles
        assert "body_indent" in styles
        assert "image" in styles
        assert "figcaption" in styles
        assert "quote" in styles
        assert "code" in styles
        assert "bullet_list" in styles
        assert "number_list" in styles

    def test_body_fallback(self, tmp_path):
        """If no body guide, a default empty format is created."""
        from docx import Document
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("一级标题")
        path = tmp_path / "no_body.docx"
        doc.save(str(path))

        styles = extract_template_styles(path)
        assert "body" in styles
        assert styles["body"].font_name is None

    def test_first_match_wins(self, tmp_path):
        from docx import Document
        from docx.shared import Pt
        doc = Document()
        p1 = doc.add_paragraph()
        p1.add_run("正文").font.size = Pt(12)
        p2 = doc.add_paragraph()
        p2.add_run("正文（另一种）").font.size = Pt(14)
        path = tmp_path / "first_wins.docx"
        doc.save(str(path))

        styles = extract_template_styles(path)
        assert styles["body"].font_size_emu == 12 * 12700


class TestValidateTemplate:
    def test_valid_template(self, minimal_template):
        result = validate_template(minimal_template)
        assert result["missing_required"] == []

    def test_missing_required(self, tmp_path):
        from docx import Document
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("正文")
        path = tmp_path / "minimal.docx"
        doc.save(str(path))

        result = validate_template(path)
        assert "h1" in result["missing_required"]
        assert "h2" in result["missing_required"]
        assert "h3" in result["missing_required"]
        assert "code" in result["missing_required"]
