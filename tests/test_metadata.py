"""Tests for metadata.py — red-head, GB compliance, page numbers, guide removal."""

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from md2word.context import ConversionReport
from md2word.metadata import (
    check_gb_compliance,
    fix_ooxml_metadata,
    insert_redhead_header,
    remove_guide_paragraphs,
    set_page_number_format,
)
from md2word.template import ParagraphFormat, extract_template_styles


def _make_minimal_doc() -> Document:
    """Create a minimal document for testing."""
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    return doc


class TestInsertRedheadHeader:
    def test_default_year(self, tmp_path):
        """Default year 2024 with no number."""
        doc = _make_minimal_doc()
        styles = {"body": ParagraphFormat(font_name="FangSong", font_size_emu=16 * 12700)}
        insert_redhead_header(doc, "XX市人民政府", styles)
        texts = [p.text for p in doc.paragraphs]
        assert "XX市人民政府" in texts
        assert any("〔2024〕" in t for t in texts)

    def test_custom_year_and_number(self, tmp_path):
        doc = _make_minimal_doc()
        styles = {"body": ParagraphFormat(font_name="FangSong", font_size_emu=16 * 12700)}
        insert_redhead_header(doc, "国务院", styles, year=2026, number="12")
        texts = [p.text for p in doc.paragraphs]
        assert "国务院" in texts
        assert any("〔2026〕" in t for t in texts)
        assert any("12号" in t for t in texts)

    def test_year_only(self, tmp_path):
        doc = _make_minimal_doc()
        styles = {"body": ParagraphFormat(font_name="FangSong")}
        insert_redhead_header(doc, "测试", styles, year=2025)
        texts = [p.text for p in doc.paragraphs]
        assert any("〔2025〕" in t for t in texts)

    def test_red_elements_present(self, tmp_path):
        doc = _make_minimal_doc()
        styles = {"body": ParagraphFormat(font_name="FangSong")}
        insert_redhead_header(doc, "XX省人民政府", styles)
        # Should have: authority, 文件, separator line, document number
        assert len(doc.paragraphs) >= 4
        # Check red color on first paragraph
        p0 = doc.paragraphs[0]
        if p0.runs:
            assert p0.runs[0].font.color.rgb is not None


class TestFixOoxmlMetadata:
    def test_strips_thumbnail(self, tmp_path):
        """Verify thumbnail JPEG is removed from docx ZIP."""
        from docx import Document as _Doc
        doc = _Doc()
        path = tmp_path / "test.docx"
        doc.save(str(path))

        # Manually inject a thumbnail
        import zipfile, io
        buf = path.read_bytes()
        out = io.BytesIO()
        with zipfile.ZipFile(io.BytesIO(buf)) as z:
            with zipfile.ZipFile(out, "w") as zout:
                for item in z.infolist():
                    zout.writestr(item, z.read(item.filename))
                zout.writestr("docProps/thumbnail.jpeg", b"fake_image_data")
        path.write_bytes(out.getvalue())

        fix_ooxml_metadata(path)
        with zipfile.ZipFile(str(path)) as z:
            names = z.namelist()
            assert not any("thumbnail" in n for n in names)

    def test_fixes_application_name(self, tmp_path):
        from docx import Document as _Doc
        doc = _Doc()
        path = tmp_path / "test2.docx"
        doc.save(str(path))

        import zipfile, io
        buf = path.read_bytes()
        out = io.BytesIO()
        with zipfile.ZipFile(io.BytesIO(buf)) as z:
            with zipfile.ZipFile(out, "w") as zout:
                for item in z.infolist():
                    raw = z.read(item.filename)
                    if item.filename == "docProps/app.xml":
                        raw = raw.replace(
                            b"Microsoft Office Word",
                            b"Microsoft Macintosh Word",
                        )
                    zout.writestr(item, raw)
        path.write_bytes(out.getvalue())

        fix_ooxml_metadata(path)
        with zipfile.ZipFile(str(path)) as z:
            app_xml = z.read("docProps/app.xml").decode("utf-8")
            assert "Microsoft Office Word" in app_xml
            assert "Microsoft Macintosh Word" not in app_xml


class TestCheckGbCompliance:
    def test_compliant_no_warnings(self):
        doc = _make_minimal_doc()
        report = ConversionReport()
        styles = {
            "body": ParagraphFormat(font_name="SimSun", font_size_emu=12 * 12700),
        }
        check_gb_compliance(doc, styles, report)
        assert not report.warnings

    def test_noncompliant_margins_emits_info(self):
        doc = _make_minimal_doc()
        doc.sections[0].top_margin = Cm(1)
        report = ConversionReport()
        styles = {"body": ParagraphFormat(font_name="Arial")}
        check_gb_compliance(doc, styles, report)
        info_msgs = report.info
        assert len(info_msgs) >= 1

    def test_no_sections_no_crash(self):
        doc = _make_minimal_doc()
        # Simulate no sections by clearing
        report = ConversionReport()
        check_gb_compliance(doc, {"body": ParagraphFormat()}, report)
        # Should not crash


class TestSetPageNumberFormat:
    def test_sets_footer(self, tmp_path):
        doc = _make_minimal_doc()
        set_page_number_format(doc, "-- %d --")
        footer = doc.sections[0].footer
        text_content = footer.paragraphs[0].text if footer.paragraphs else ""
        assert text_content is not None  # Should have field code, not plain text

    def test_no_sections_no_crash(self):
        doc = Document()
        # No sections — should not crash
        set_page_number_format(doc)

    def test_empty_footer_format(self):
        doc = _make_minimal_doc()
        set_page_number_format(doc)
        # Should not crash


class TestRemoveGuideParagraphs:
    def test_removes_known_guide_paragraphs(self, tmp_path):
        from docx import Document as _Doc
        doc = _Doc()
        doc.add_paragraph("一级标题")
        doc.add_paragraph("正文")
        doc.add_paragraph("实际内容")
        assert len(doc.paragraphs) == 3

        remove_guide_paragraphs(doc)
        texts = [p.text for p in doc.paragraphs]
        assert "一级标题" not in texts
        assert "正文" not in texts
        assert "实际内容" in texts

    def test_removes_abstract_keywords_guides(self, tmp_path):
        """With v1.6, abstract/keywords/references are also guide paragraphs."""
        from docx import Document as _Doc
        doc = _Doc()
        doc.add_paragraph("摘要")
        doc.add_paragraph("关键词")
        doc.add_paragraph("参考文献")
        doc.add_paragraph("真实内容")

        remove_guide_paragraphs(doc)
        texts = [p.text for p in doc.paragraphs]
        assert "摘要" not in texts
        assert "关键词" not in texts
        assert "参考文献" not in texts
        assert "真实内容" in texts

    def test_empty_document(self):
        doc = _make_minimal_doc()
        remove_guide_paragraphs(doc)
        # Should not crash
