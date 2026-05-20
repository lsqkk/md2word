"""Shared test fixtures."""

from pathlib import Path

import pytest


@pytest.fixture
def sample_markdown() -> str:
    return """# 一级标题

这是一段正文内容。

## 二级标题

- 无序列表项1
- 无序列表项2

1. 有序列表项A
2. 有序列表项B

> 这是一段引用

```python
def hello():
    print("Hello")
```

| 列1 | 列2 |
|-----|-----|
| A   | B   |

这是一段有脚注的文字[^1].

[^1]: 脚注内容
"""


@pytest.fixture
def markdown_with_footnotes() -> str:
    return """正文内容[^1]和更多内容[^2].

[^1]: 第一个脚注
[^2]: 第二个脚注，多行
    继续
"""


@pytest.fixture
def markdown_with_math() -> str:
    return """行内公式 $E=mc^2$ 和块公式:

$$\\int_0^\\infty e^{-x^2} dx = \\frac{\\sqrt{\\pi}}{2}$$
"""


@pytest.fixture
def markdown_with_front_matter() -> str:
    return """---
title: 测试文档
author: Me
date: 2024-01-01
---

# 正文开始
"""


@pytest.fixture
def markdown_without_lists() -> str:
    """Markdown where list markers need blank lines inserted."""
    return "前面有文字\n- 列表项1\n- 列表项2\n\n1. 有序项"


@pytest.fixture
def minimal_template(tmp_path: Path) -> Path:
    """Create a minimal .docx template with guide paragraphs."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    # Set margins
    section = doc.sections[0]
    from docx.shared import Cm
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # Guide paragraphs
    guide_data = [
        ("一级标题", "SimHei", 18, True),
        ("二级标题", "SimHei", 15, True),
        ("三级标题", "SimHei", 13, True),
        ("正文", "SimSun", 12, False),
        ("首行缩进", "SimSun", 12, False),
        ("图片", "SimSun", 10, False),
        ("图注", "SimSun", 9, False),
        ("引用", "KaiTi", 10, False),
        ("代码", "DengXian", 9, False),
        ("无序列表", "SimSun", 12, False),
        ("有序列表", "SimSun", 12, False),
    ]

    for text, font, size, bold in guide_data:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = font
        run.font.size = Pt(size)
        run.font.bold = bold

    path = tmp_path / "minimal_template.docx"
    doc.save(str(path))
    return path
