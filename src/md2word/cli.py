"""CLI entry point for md2word."""

from __future__ import annotations

import argparse
import sys
from pathlib import Path

from . import __version__
from .converter import convert
from .template import list_template_styles

# ── template definitions ────────────────────────────────────────────────────

_THEMES: dict[str, dict] = {}


def _register_theme(name: str, label: str, desc: str, filename: str, builder):
    _THEMES[name] = {
        "label": label,
        "desc": desc,
        "filename": filename,
        "builder": builder,
    }


def _set_run_font_ea(run, western: str, eastern: str | None = None):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), western)
    rFonts.set(qn("w:hAnsi"), western)
    rFonts.set(qn("w:eastAsia"), eastern or western)


# ── Theme 1: 官方公文 ─────────────────────────────────────────────────────

def _build_official(output_path: Path):
    """官方公文风 — 仿照《党政机关公文格式》GB/T 9704-2012."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    doc = Document()
    s = doc.sections[0]
    s.top_margin = Cm(3.7)
    s.bottom_margin = Cm(3.5)
    s.left_margin = Cm(2.8)
    s.right_margin = Cm(2.6)

    # 一级标题 — 黑体 二号 居中
    p = doc.add_paragraph()
    r = p.add_run("一级标题")
    r.font.size = Pt(22)
    r.font.bold = True
    _set_run_font_ea(r, "SimHei", "SimHei")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)

    # 二级标题 — 黑体 三号
    p = doc.add_paragraph()
    r = p.add_run("二级标题")
    r.font.size = Pt(16)
    r.font.bold = True
    _set_run_font_ea(r, "SimHei", "SimHei")
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)

    # 三级标题 — 楷体 三号
    p = doc.add_paragraph()
    r = p.add_run("三级标题")
    r.font.size = Pt(16)
    r.font.bold = True
    _set_run_font_ea(r, "KaiTi", "KaiTi")
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    # 正文 — 仿宋 三号
    p = doc.add_paragraph()
    r = p.add_run("正文")
    r.font.size = Pt(16)
    _set_run_font_ea(r, "FangSong", "FangSong")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5

    # 首行缩进 — 仿宋 三号 首行缩进2字符
    p = doc.add_paragraph()
    r = p.add_run("首行缩进")
    r.font.size = Pt(16)
    _set_run_font_ea(r, "FangSong", "FangSong")
    p.paragraph_format.first_line_indent = Cm(0.85)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5

    # 图片 — 仿宋 五号 居中
    p = doc.add_paragraph()
    r = p.add_run("图片")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    _set_run_font_ea(r, "FangSong", "FangSong")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 图注 — 仿宋 小五 居中
    p = doc.add_paragraph()
    r = p.add_run("图注")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    _set_run_font_ea(r, "FangSong", "FangSong")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 引用 — 楷体 四号
    p = doc.add_paragraph()
    r = p.add_run("引用")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    _set_run_font_ea(r, "KaiTi", "KaiTi")
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    # 代码 — 等线 五号
    p = doc.add_paragraph()
    r = p.add_run("代码")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    _set_run_font_ea(r, "DengXian", "DengXian")
    p.paragraph_format.left_indent = Cm(0.5)

    # 无序列表 — 仿宋 三号
    p = doc.add_paragraph()
    r = p.add_run("无序列表")
    r.font.size = Pt(16)
    _set_run_font_ea(r, "FangSong", "FangSong")
    p.paragraph_format.left_indent = Cm(0.75)

    # 有序列表 — 仿宋 三号
    p = doc.add_paragraph()
    r = p.add_run("有序列表")
    r.font.size = Pt(16)
    _set_run_font_ea(r, "FangSong", "FangSong")
    p.paragraph_format.left_indent = Cm(0.75)

    doc.save(str(output_path))


_register_theme("official", "官方公文", "仿公文GB标准：仿宋正文+黑体标题+大字号", "官方公文.docx", _build_official)


# ── Theme 2: 学术论文 ─────────────────────────────────────────────────────


def _build_academic(output_path: Path):
    """学术论文风 — 全宋体系列、标准学术格式."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    doc = Document()
    s = doc.sections[0]
    s.top_margin = Cm(2.54)
    s.bottom_margin = Cm(2.54)
    s.left_margin = Cm(3.17)
    s.right_margin = Cm(3.17)

    # 一级标题 — 黑体 三号 居中
    p = doc.add_paragraph()
    r = p.add_run("一级标题")
    r.font.size = Pt(16)
    r.font.bold = True
    _set_run_font_ea(r, "SimHei", "SimHei")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)

    # 二级标题 — 黑体 四号
    p = doc.add_paragraph()
    r = p.add_run("二级标题")
    r.font.size = Pt(14)
    r.font.bold = True
    _set_run_font_ea(r, "SimHei", "SimHei")
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)

    # 三级标题 — 黑体 小四
    p = doc.add_paragraph()
    r = p.add_run("三级标题")
    r.font.size = Pt(12)
    r.font.bold = True
    _set_run_font_ea(r, "SimHei", "SimHei")
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)

    # 正文 — 宋体 小四
    p = doc.add_paragraph()
    r = p.add_run("正文")
    r.font.size = Pt(12)
    _set_run_font_ea(r, "SimSun", "SimSun")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5

    # 首行缩进 — 宋体 小四 首行缩进2字符
    p = doc.add_paragraph()
    r = p.add_run("首行缩进")
    r.font.size = Pt(12)
    _set_run_font_ea(r, "SimSun", "SimSun")
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5

    # 图片 — 宋体 五号 居中
    p = doc.add_paragraph()
    r = p.add_run("图片")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    _set_run_font_ea(r, "SimSun", "SimSun")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)

    # 图注 — 宋体 小五 居中
    p = doc.add_paragraph()
    r = p.add_run("图注")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    _set_run_font_ea(r, "SimSun", "SimSun")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 引用 — 楷体 五号
    p = doc.add_paragraph()
    r = p.add_run("引用")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    _set_run_font_ea(r, "KaiTi", "KaiTi")
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)

    # 代码 — 宋体 小五 （论文一般不特别区分代码）
    p = doc.add_paragraph()
    r = p.add_run("代码")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    _set_run_font_ea(r, "DengXian", "DengXian")
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

    # 无序列表 — 宋体 小四
    p = doc.add_paragraph()
    r = p.add_run("无序列表")
    r.font.size = Pt(12)
    _set_run_font_ea(r, "SimSun", "SimSun")
    p.paragraph_format.left_indent = Cm(0.75)

    # 有序列表 — 宋体 小四
    p = doc.add_paragraph()
    r = p.add_run("有序列表")
    r.font.size = Pt(12)
    _set_run_font_ea(r, "SimSun", "SimSun")
    p.paragraph_format.left_indent = Cm(0.75)

    doc.save(str(output_path))


_register_theme("academic", "学术论文", "全宋体系列+黑体标题+标准学术排版", "学术论文.docx", _build_academic)


# ── Theme 3: 技术文档 ─────────────────────────────────────────────────────


def _build_tech(output_path: Path):
    """技术文档风 — 现代、代码友好、信息型."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    doc = Document()
    s = doc.sections[0]
    s.top_margin = Cm(2)
    s.bottom_margin = Cm(2)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)

    # 一级标题 — 微软雅黑 18pt 深蓝
    p = doc.add_paragraph()
    r = p.add_run("一级标题")
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    _set_run_font_ea(r, "Microsoft YaHei", "Microsoft YaHei")
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(8)

    # 二级标题 — 微软雅黑 14pt 深灰
    p = doc.add_paragraph()
    r = p.add_run("二级标题")
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    _set_run_font_ea(r, "Microsoft YaHei", "Microsoft YaHei")
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)

    # 三级标题 — 微软雅黑 12pt
    p = doc.add_paragraph()
    r = p.add_run("三级标题")
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    _set_run_font_ea(r, "Microsoft YaHei", "Microsoft YaHei")
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)

    # 正文 — 微软雅黑 10.5pt 深灰
    p = doc.add_paragraph()
    r = p.add_run("正文")
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    _set_run_font_ea(r, "Microsoft YaHei", "Microsoft YaHei")
    p.paragraph_format.line_spacing = 1.3

    # 首行缩进
    p = doc.add_paragraph()
    r = p.add_run("首行缩进")
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    _set_run_font_ea(r, "Microsoft YaHei", "Microsoft YaHei")
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.3

    # 图片
    p = doc.add_paragraph()
    r = p.add_run("图片")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    _set_run_font_ea(r, "Microsoft YaHei", "Microsoft YaHei")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)

    # 图注
    p = doc.add_paragraph()
    r = p.add_run("图注")
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    _set_run_font_ea(r, "Microsoft YaHei", "Microsoft YaHei")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 引用 — 左边框线
    p = doc.add_paragraph()
    r = p.add_run("引用")
    r.font.size = Pt(10.5)
    r.font.italic = False
    r.font.color.rgb = RGBColor(0x55, 0x6B, 0x82)
    _set_run_font_ea(r, "Microsoft YaHei", "Microsoft YaHei")
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

    # 代码 — Consolas + 等线
    p = doc.add_paragraph()
    r = p.add_run("代码")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    _set_run_font_ea(r, "Consolas", "DengXian")
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)

    # 无序列表
    p = doc.add_paragraph()
    r = p.add_run("无序列表")
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    _set_run_font_ea(r, "Microsoft YaHei", "Microsoft YaHei")
    p.paragraph_format.left_indent = Cm(0.75)

    # 有序列表
    p = doc.add_paragraph()
    r = p.add_run("有序列表")
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    _set_run_font_ea(r, "Microsoft YaHei", "Microsoft YaHei")
    p.paragraph_format.left_indent = Cm(0.75)

    doc.save(str(output_path))


_register_theme("tech", "技术文档", "现代微软雅黑+深蓝色标题+紧凑排版", "技术文档.docx", _build_tech)


# ── Theme 4: 自媒体排版 ──────────────────────────────────────────────────


def _build_media(output_path: Path):
    """自媒体排版风 — 时尚视觉、高行距、大字报风格."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    doc = Document()
    s = doc.sections[0]
    s.top_margin = Cm(1.5)
    s.bottom_margin = Cm(1.5)
    s.left_margin = Cm(2)
    s.right_margin = Cm(2)

    # 一级标题 — 微软雅黑 24pt 品牌橙
    p = doc.add_paragraph()
    r = p.add_run("一级标题")
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0xE6, 0x4A, 0x19)
    _set_run_font_ea(r, "Microsoft YaHei", "Microsoft YaHei")
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(14)

    # 二级标题 — 微软雅黑 18pt 品牌橙
    p = doc.add_paragraph()
    r = p.add_run("二级标题")
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0xE6, 0x4A, 0x19)
    _set_run_font_ea(r, "Microsoft YaHei", "Microsoft YaHei")
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(8)

    # 三级标题 — 微软雅黑 15pt 深灰
    p = doc.add_paragraph()
    r = p.add_run("三级标题")
    r.font.size = Pt(15)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    _set_run_font_ea(r, "Microsoft YaHei", "Microsoft YaHei")
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)

    # 正文 — 思源宋体/宋体 12pt 高行距
    p = doc.add_paragraph()
    r = p.add_run("正文")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x3A, 0x3A, 0x3A)
    _set_run_font_ea(r, "SimSun", "SimSun")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.8

    # 首行缩进
    p = doc.add_paragraph()
    r = p.add_run("首行缩进")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x3A, 0x3A, 0x3A)
    _set_run_font_ea(r, "SimSun", "SimSun")
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.8

    # 图片 — 居中宽松
    p = doc.add_paragraph()
    r = p.add_run("图片")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
    r.font.italic = True
    _set_run_font_ea(r, "Microsoft YaHei", "Microsoft YaHei")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)

    # 图注
    p = doc.add_paragraph()
    r = p.add_run("图注")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
    _set_run_font_ea(r, "Microsoft YaHei", "Microsoft YaHei")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 引用 — 楷体 大字号 浅灰
    p = doc.add_paragraph()
    r = p.add_run("引用")
    r.font.size = Pt(14)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    _set_run_font_ea(r, "KaiTi", "KaiTi")
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)

    # 代码 — 等宽
    p = doc.add_paragraph()
    r = p.add_run("代码")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    _set_run_font_ea(r, "Consolas", "DengXian")
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

    # 无序列表
    p = doc.add_paragraph()
    r = p.add_run("无序列表")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x3A, 0x3A, 0x3A)
    _set_run_font_ea(r, "SimSun", "SimSun")
    p.paragraph_format.left_indent = Cm(0.75)

    # 有序列表
    p = doc.add_paragraph()
    r = p.add_run("有序列表")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x3A, 0x3A, 0x3A)
    _set_run_font_ea(r, "SimSun", "SimSun")
    p.paragraph_format.left_indent = Cm(0.75)

    doc.save(str(output_path))


_register_theme("media", "自媒体排版", "视觉系大字报风格+高行距+品牌色", "自媒体排版.docx", _build_media)


# ── CLI ──────────────────────────────────────────────────────────────────────


def _create_template(output_path: Path, theme: str = "official") -> None:
    if theme not in _THEMES:
        print(f"Unknown theme '{theme}'. Available: {', '.join(_THEMES)}")
        print("Falling back to 'official'.")
        theme = "official"
    _THEMES[theme]["builder"](output_path)
    info = _THEMES[theme]
    print(f"Template created: {output_path}")
    print(f"  Theme: {info['label']} — {info['desc']}")
    print("Open it in Word, adjust any style, then save.")
    print("The tool detects styles by marker keywords (一级标题, 正文, etc.).")


def parse_args(argv: list[str] | None = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        prog="md2word",
        description="Convert Markdown to Word (.docx) using a custom template.",
    )

    parser.add_argument(
        "input", nargs="?", type=str, help="Input Markdown file (omit to read stdin)"
    )
    parser.add_argument("-o", "--output", type=str, default=None, help="Output .docx path")
    parser.add_argument(
        "-t", "--template", type=str, default=None, help="Template .docx with guide paragraphs"
    )
    parser.add_argument(
        "--image-width", type=float, default=5.5, help="Max image width in inches (default: 5.5)"
    )
    parser.add_argument(
        "--list-styles", action="store_true", help="List detected guide styles in a template"
    )
    parser.add_argument(
        "--create-template",
        type=str,
        default=None,
        metavar="PATH",
        help="Generate a sample template (use --theme to choose style)",
    )
    parser.add_argument(
        "--theme",
        type=str,
        default="official",
        choices=list(_THEMES),
        help="Template theme (use --list-themes to see all)",
    )
    parser.add_argument(
        "--list-themes", action="store_true", help="List available template themes"
    )
    parser.add_argument("--version", action="store_true", help="Show version")

    return parser.parse_args(argv)


def main(argv: list[str] | None = None) -> int:
    if sys.stdout.encoding and sys.stdout.encoding.upper() in ("GBK", "GB2312", "GB18030"):
        sys.stdout.reconfigure(errors="replace")
    args = parse_args(argv)

    if args.version:
        print(f"md2word v{__version__}")
        return 0

    if args.list_themes:
        print("Available template themes:\n")
        for name, info in _THEMES.items():
            print(f"  {name:<12} {info['label']:<14} {info['filename']:<16}  {info['desc']}")
        return 0

    if args.create_template:
        _create_template(Path(args.create_template), theme=args.theme)
        return 0

    if args.list_styles:
        tpl = args.template or _find_default_template()
        if not tpl:
            print("No template specified and no default found.", file=sys.stderr)
            return 1
        items = list_template_styles(tpl)
        if not items:
            print(f"No guide paragraphs found in {tpl}")
            return 0
        print(f"Styles in: {tpl}\n")
        print(f"{'Slot':<14} {'Guide Text':<12} {'Font':<18} {'EA Font':<14} {'Size':<7} {'Bold':<6}")
        print("-" * 80)
        for item in items:
            print(
                f"{item['slot']:<14} {item['guide_text'][:10]:<12} "
                f"{item['font']:<18} {item['ea_font']:<14} "
                f"{str(item['size_pt']):<7} {'Yes' if item['bold'] else 'No':<6}"
            )
        return 0

    # Input
    if args.input:
        in_path = Path(args.input)
        if not in_path.exists():
            print(f"Input file not found: {in_path}", file=sys.stderr)
            return 1
        md_text = in_path.read_text(encoding="utf-8")
    else:
        md_text = sys.stdin.read()
        in_path = None

    # Output
    if args.output:
        out_path = Path(args.output)
    elif in_path:
        out_path = in_path.with_suffix(".docx")
    else:
        print("Output path required when reading from stdin (use -o)", file=sys.stderr)
        return 1

    # Template
    tpl = args.template or _find_default_template()
    if not tpl:
        print("No template specified and no default found.", file=sys.stderr)
        print("Generate one: md2word --create-template template.docx")
        return 1

    print(f"Converting: {args.input or '(stdin)'}")
    print(f"Template:   {tpl}")
    print(f"Output:     {out_path}")

    convert(md_text, tpl, out_path, image_max_width=args.image_width)
    print(f"Done → {out_path}")
    return 0


def _find_default_template() -> Path | None:
    """Find default template — prefers theme-named files in template/."""
    search = [
        Path("template/template1.docx"),
        Path("template/官方公文.docx"),
        Path("template/学术论文.docx"),
        Path("template/技术文档.docx"),
        Path("template/自媒体排版.docx"),
    ]
    # Also look relative to package
    pkg_tpl = Path(__file__).parent.parent.parent / "template"
    search.extend(pkg_tpl / n for n in [
        "template1.docx", "官方公文.docx", "学术论文.docx",
        "技术文档.docx", "自媒体排版.docx",
    ])
    for c in search:
        if c.exists():
            return c
    return None


if __name__ == "__main__":
    sys.exit(main())
