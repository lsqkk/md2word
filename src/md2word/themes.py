"""Theme specifications for template generation — data-driven, zero boilerplate."""
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from .template import _inject_slot_marker, set_run_font

_SLOT_NAMES: dict[str, str] = {
    "一级标题": "h1",
    "二级标题": "h2",
    "三级标题": "h3",
    "正文": "body",
    "首行缩进": "body_indent",
    "图片": "image",
    "图注": "figcaption",
    "引用": "quote",
    "代码": "code",
    "无序列表": "bullet_list",
    "有序列表": "number_list",
    "目录标题": "toc_title",
    "摘要": "abstract",
    "关键词": "keywords",
    "参考文献": "references",
}


@dataclass
class ParagraphSpec:
    """Spec for a single guide paragraph in a theme template."""
    label: str
    font: str
    ea_font: str | None = None
    size_pt: float = 12
    bold: bool = False
    italic: bool = False
    color: tuple[int, int, int] | None = None
    align: int | None = None          # WD_ALIGN_PARAGRAPH value
    first_line_indent_cm: float | None = None
    left_indent_cm: float | None = None
    space_before_pt: float = 0
    space_after_pt: float = 0
    line_spacing: float | None = None


@dataclass
class ThemeSpec:
    """Complete specification for a theme template."""
    label: str
    filename: str
    desc: str
    margins_cm: tuple[float, float, float, float]  # top, bottom, left, right
    paragraphs: list[ParagraphSpec]


# ── Theme definitions ──────────────────────────────────────────────────────────

OFFICIAL = ThemeSpec(
    label="官方公文",
    filename="官方公文.docx",
    desc="仿公文GB标准：仿宋正文+黑体标题+大字号",
    margins_cm=(3.7, 3.5, 2.8, 2.6),
    paragraphs=[
        ParagraphSpec("一级标题", "SimHei", size_pt=22, bold=True,
                       align=1, space_before_pt=24, space_after_pt=12),
        ParagraphSpec("二级标题", "SimHei", size_pt=16, bold=True,
                       space_before_pt=12, space_after_pt=6),
        ParagraphSpec("三级标题", "KaiTi", size_pt=16, bold=True,
                       space_before_pt=6, space_after_pt=6),
        ParagraphSpec("正文", "FangSong", size_pt=16,
                       align=3, space_after_pt=6, line_spacing=1.5),
        ParagraphSpec("首行缩进", "FangSong", size_pt=16,
                       first_line_indent_cm=0.85, align=3, line_spacing=1.5),
        ParagraphSpec("图片", "FangSong", size_pt=10,
                       color=(0x88, 0x88, 0x88), align=1),
        ParagraphSpec("图注", "FangSong", size_pt=9,
                       color=(0x88, 0x88, 0x88), align=1),
        ParagraphSpec("引用", "KaiTi", size_pt=14,
                       color=(0x55, 0x55, 0x55),
                       left_indent_cm=1, space_before_pt=6, space_after_pt=6),
        ParagraphSpec("代码", "DengXian", size_pt=10,
                       color=(0x33, 0x33, 0x33), left_indent_cm=0.5),
        ParagraphSpec("无序列表", "FangSong", size_pt=16, left_indent_cm=0.75),
        ParagraphSpec("有序列表", "FangSong", size_pt=16, left_indent_cm=0.75),
    ],
)

ACADEMIC = ThemeSpec(
    label="学术论文",
    filename="学术论文.docx",
    desc="全宋体系列+黑体标题居中+标准学术排版+参考文献格式",
    margins_cm=(2.54, 2.54, 3.17, 3.17),
    paragraphs=[
        ParagraphSpec("一级标题", "SimHei", size_pt=16, bold=True,
                       align=1, space_before_pt=18, space_after_pt=12),
        ParagraphSpec("二级标题", "SimHei", size_pt=14, bold=True,
                       space_before_pt=12, space_after_pt=6),
        ParagraphSpec("三级标题", "SimHei", size_pt=12, bold=True,
                       space_before_pt=6, space_after_pt=3),
        ParagraphSpec("正文", "SimSun", size_pt=12,
                       align=3, line_spacing=1.5),
        ParagraphSpec("首行缩进", "SimSun", size_pt=12,
                       first_line_indent_cm=0.74, align=3, line_spacing=1.5),
        ParagraphSpec("图片", "SimSun", size_pt=10,
                       color=(0x55, 0x55, 0x55), align=1, space_before_pt=6),
        ParagraphSpec("图注", "SimSun", size_pt=9,
                       color=(0x66, 0x66, 0x66), align=1),
        ParagraphSpec("引用", "KaiTi", size_pt=10,
                       color=(0x44, 0x44, 0x44),
                       left_indent_cm=1.5, space_before_pt=3, space_after_pt=3),
        ParagraphSpec("代码", "DengXian", size_pt=9,
                       color=(0x33, 0x33, 0x33),
                       left_indent_cm=0.5, space_before_pt=2, space_after_pt=2),
        ParagraphSpec("无序列表", "SimSun", size_pt=12, left_indent_cm=0.75),
        ParagraphSpec("有序列表", "SimSun", size_pt=12, left_indent_cm=0.75),
    ],
)

ACADEMIC_ENHANCED = ThemeSpec(
    label="学术论文（增强版）",
    filename="学术论文_增强版.docx",
    desc="GB/T 7713 标准：摘要+关键词+参考文献+页眉页脚+三线表",
    margins_cm=(2.54, 2.54, 3.17, 3.17),
    paragraphs=[
        ParagraphSpec("一级标题", "SimHei", size_pt=16, bold=True,
                       align=1, space_before_pt=18, space_after_pt=12),
        ParagraphSpec("二级标题", "SimHei", size_pt=14, bold=True,
                       space_before_pt=12, space_after_pt=6),
        ParagraphSpec("三级标题", "SimHei", size_pt=12, bold=True,
                       space_before_pt=6, space_after_pt=3),
        ParagraphSpec("正文", "SimSun", size_pt=12,
                       align=3, line_spacing=1.5),
        ParagraphSpec("首行缩进", "SimSun", size_pt=12,
                       first_line_indent_cm=0.74, align=3, line_spacing=1.5),
        ParagraphSpec("摘要", "SimHei", size_pt=12, bold=True,
                       space_before_pt=12),
        ParagraphSpec("关键词", "SimHei", size_pt=12, bold=True,
                       space_before_pt=6, space_after_pt=12),
        ParagraphSpec("图片", "SimSun", size_pt=10,
                       color=(0x55, 0x55, 0x55), align=1, space_before_pt=6),
        ParagraphSpec("图注", "SimSun", size_pt=9,
                       color=(0x66, 0x66, 0x66), align=1),
        ParagraphSpec("引用", "KaiTi", size_pt=10,
                       color=(0x44, 0x44, 0x44),
                       left_indent_cm=1.5, space_before_pt=3, space_after_pt=3),
        ParagraphSpec("代码", "DengXian", size_pt=9,
                       color=(0x33, 0x33, 0x33),
                       left_indent_cm=0.5, space_before_pt=2, space_after_pt=2),
        ParagraphSpec("无序列表", "SimSun", size_pt=12, left_indent_cm=0.75),
        ParagraphSpec("有序列表", "SimSun", size_pt=12, left_indent_cm=0.75),
        ParagraphSpec("参考文献", "SimSun", size_pt=9,
                       line_spacing=1.25),
    ],
)

TECH = ThemeSpec(
    label="技术文档",
    filename="技术文档.docx",
    desc="现代微软雅黑+深蓝色标题+紧凑排版+代码友好",
    margins_cm=(2, 2, 2.5, 2.5),
    paragraphs=[
        ParagraphSpec("一级标题", "Microsoft YaHei", size_pt=18, bold=True,
                       color=(0x1A, 0x3C, 0x6E),
                       space_before_pt=20, space_after_pt=8),
        ParagraphSpec("二级标题", "Microsoft YaHei", size_pt=14, bold=True,
                       color=(0x33, 0x33, 0x33),
                       space_before_pt=14, space_after_pt=6),
        ParagraphSpec("三级标题", "Microsoft YaHei", size_pt=12, bold=True,
                       color=(0x44, 0x44, 0x44),
                       space_before_pt=10, space_after_pt=4),
        ParagraphSpec("正文", "Microsoft YaHei", size_pt=10.5,
                       color=(0x33, 0x33, 0x33), line_spacing=1.3),
        ParagraphSpec("首行缩进", "Microsoft YaHei", size_pt=10.5,
                       color=(0x33, 0x33, 0x33),
                       first_line_indent_cm=0.74, line_spacing=1.3),
        ParagraphSpec("图片", "Microsoft YaHei", size_pt=9,
                       color=(0x88, 0x88, 0x88), align=1, space_before_pt=8),
        ParagraphSpec("图注", "Microsoft YaHei", size_pt=8,
                       color=(0xAA, 0xAA, 0xAA), align=1),
        ParagraphSpec("引用", "Microsoft YaHei", size_pt=10.5,
                       color=(0x55, 0x6B, 0x82),
                       left_indent_cm=1, space_before_pt=4, space_after_pt=4),
        ParagraphSpec("代码", "Consolas", ea_font="DengXian", size_pt=9,
                       color=(0x1A, 0x1A, 0x2E),
                       left_indent_cm=0.5, space_before_pt=3, space_after_pt=3),
        ParagraphSpec("无序列表", "Microsoft YaHei", size_pt=10.5,
                       color=(0x33, 0x33, 0x33), left_indent_cm=0.75),
        ParagraphSpec("有序列表", "Microsoft YaHei", size_pt=10.5,
                       color=(0x33, 0x33, 0x33), left_indent_cm=0.75),
    ],
)

MEDIA = ThemeSpec(
    label="自媒体排版",
    filename="自媒体排版.docx",
    desc="视觉系大字报风格+高行距+品牌色+社交传播",
    margins_cm=(1.5, 1.5, 2, 2),
    paragraphs=[
        ParagraphSpec("一级标题", "Microsoft YaHei", size_pt=24, bold=True,
                       color=(0xE6, 0x4A, 0x19),
                       space_before_pt=28, space_after_pt=14),
        ParagraphSpec("二级标题", "Microsoft YaHei", size_pt=18, bold=True,
                       color=(0xE6, 0x4A, 0x19),
                       space_before_pt=20, space_after_pt=8),
        ParagraphSpec("三级标题", "Microsoft YaHei", size_pt=15, bold=True,
                       color=(0x33, 0x33, 0x33),
                       space_before_pt=14, space_after_pt=6),
        ParagraphSpec("正文", "SimSun", size_pt=12,
                       color=(0x3A, 0x3A, 0x3A),
                       align=3, space_after_pt=12, line_spacing=1.8),
        ParagraphSpec("首行缩进", "SimSun", size_pt=12,
                       color=(0x3A, 0x3A, 0x3A),
                       first_line_indent_cm=0.74, align=3, line_spacing=1.8),
        ParagraphSpec("图片", "Microsoft YaHei", size_pt=10,
                       color=(0xBB, 0xBB, 0xBB), italic=True,
                       align=1, space_before_pt=12),
        ParagraphSpec("图注", "Microsoft YaHei", size_pt=10,
                       color=(0xBB, 0xBB, 0xBB), align=1),
        ParagraphSpec("引用", "KaiTi", size_pt=14,
                       color=(0x88, 0x88, 0x88), italic=True,
                       left_indent_cm=1.5, space_before_pt=12, space_after_pt=12),
        ParagraphSpec("代码", "Consolas", ea_font="DengXian", size_pt=9,
                       color=(0x1A, 0x1A, 0x1A),
                       left_indent_cm=0.5, space_before_pt=4, space_after_pt=4),
        ParagraphSpec("无序列表", "SimSun", size_pt=12,
                       color=(0x3A, 0x3A, 0x3A), left_indent_cm=0.75),
        ParagraphSpec("有序列表", "SimSun", size_pt=12,
                       color=(0x3A, 0x3A, 0x3A), left_indent_cm=0.75),
    ],
)

REDHEAD = ThemeSpec(
    label="红头文件",
    filename="红头文件.docx",
    desc="GB/T 9704-2012 党政机关公文格式标准：红头+分隔线+仿宋正文",
    margins_cm=(3.7, 3.5, 2.8, 2.6),
    paragraphs=[
        ParagraphSpec("一级标题", "方正小标宋简体", size_pt=26,
                       align=1, space_before_pt=24, space_after_pt=12),
        ParagraphSpec("二级标题", "SimHei", size_pt=16, bold=True,
                       space_before_pt=12, space_after_pt=6),
        ParagraphSpec("三级标题", "KaiTi", size_pt=16, bold=True,
                       space_before_pt=6, space_after_pt=6),
        ParagraphSpec("正文", "FangSong", size_pt=16,
                       first_line_indent_cm=0.85, align=3, space_after_pt=6, line_spacing=1.5),
        ParagraphSpec("首行缩进", "FangSong", size_pt=16,
                       first_line_indent_cm=0.85, align=3, line_spacing=1.5),
        ParagraphSpec("图片", "FangSong", size_pt=10,
                       color=(0x88, 0x88, 0x88), align=1),
        ParagraphSpec("图注", "FangSong", size_pt=9,
                       color=(0x88, 0x88, 0x88), align=1),
        ParagraphSpec("引用", "KaiTi", size_pt=14,
                       color=(0x55, 0x55, 0x55),
                       left_indent_cm=1, space_before_pt=6, space_after_pt=6),
        ParagraphSpec("代码", "DengXian", size_pt=10,
                       color=(0x33, 0x33, 0x33), left_indent_cm=0.5),
        ParagraphSpec("无序列表", "FangSong", size_pt=16, left_indent_cm=0.75),
        ParagraphSpec("有序列表", "FangSong", size_pt=16, left_indent_cm=0.75),
    ],
)

# Registry for lookup by name
_THEMES: dict[str, ThemeSpec] = {
    "official": OFFICIAL,
    "academic": ACADEMIC,
    "academic-plus": ACADEMIC_ENHANCED,
    "tech": TECH,
    "media": MEDIA,
    "redhead": REDHEAD,
}


def get_theme(name: str) -> ThemeSpec | None:
    return _THEMES.get(name)


def list_themes() -> list[tuple[str, ThemeSpec]]:
    return [(name, spec) for name, spec in _THEMES.items()]


def build_theme(spec: ThemeSpec, output_path: Path) -> None:
    """Build a template docx from a theme spec, injecting slot markers."""
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor

    doc = Document()
    s = doc.sections[0]
    s.top_margin = Cm(spec.margins_cm[0])
    s.bottom_margin = Cm(spec.margins_cm[1])
    s.left_margin = Cm(spec.margins_cm[2])
    s.right_margin = Cm(spec.margins_cm[3])
    s.start_type = 0

    for ps in spec.paragraphs:
        p = doc.add_paragraph()
        r = p.add_run(ps.label)
        r.font.size = Pt(ps.size_pt)
        r.font.bold = ps.bold
        r.font.italic = ps.italic
        if ps.color:
            r.font.color.rgb = RGBColor(*ps.color)
        set_run_font(r, ps.font, ps.ea_font or ps.font)

        if ps.align is not None:
            p.paragraph_format.alignment = ps.align
        if ps.first_line_indent_cm is not None:
            p.paragraph_format.first_line_indent = Cm(ps.first_line_indent_cm)
        if ps.left_indent_cm is not None:
            p.paragraph_format.left_indent = Cm(ps.left_indent_cm)
        if ps.space_before_pt:
            p.paragraph_format.space_before = Pt(ps.space_before_pt)
        if ps.space_after_pt:
            p.paragraph_format.space_after = Pt(ps.space_after_pt)
        if ps.line_spacing is not None:
            p.paragraph_format.line_spacing = ps.line_spacing

        # Inject machine-readable slot marker
        slot = _SLOT_NAMES.get(ps.label)
        if slot:
            _inject_slot_marker(p, slot)

    doc.save(str(output_path))
