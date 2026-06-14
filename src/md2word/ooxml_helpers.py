"""OOXML-level helpers — bookmarks, fields, SVG, numbering, borders, TOC."""

from __future__ import annotations

from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.package import ImagePart
from docx.shared import Inches, Pt, RGBColor

from .context import ConversionContext
from .image_utils import get_image_dimensions, get_svg_size, is_svg_source, resolve_image, resolve_svg_raw
from .template import ParagraphFormat


# ── Bookmark helpers ────────────────────────────────────────────────────


def add_bookmark(paragraph, ctx: ConversionContext, name: str | None = None) -> str:
    """Add a bookmark (start + end) around a paragraph.

    Returns the unique bookmark name.
    """
    bid = ctx.next_bookmark_id()
    bm_name = name or f"_Ref{bid}"
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bid))
    start.set(qn("w:name"), bm_name)
    paragraph._p.insert(0, start)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bid))
    paragraph._p.append(end)
    return bm_name


def make_ref_field(display_text: str, bookmark_name: str) -> list:
    """Create OMML runs for a REF field pointing to a bookmark.

    Returns a list of OxmlElement runs.
    """
    runs = []
    r1 = OxmlElement("w:r")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    r1.append(fld_begin)
    runs.append(r1)

    r2 = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f' REF {bookmark_name} \\h '
    r2.append(instr)
    runs.append(r2)

    r3 = OxmlElement("w:r")
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    r3.append(fld_sep)
    runs.append(r3)

    r4 = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = display_text
    r4.append(t)
    runs.append(r4)

    r5 = OxmlElement("w:r")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r5.append(fld_end)
    runs.append(r5)
    return runs


def add_seq_number(p, seq_name: str = "Equation") -> None:
    """Add a SEQ field to a paragraph for auto-numbering."""
    r1 = OxmlElement("w:r")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    r1.append(fld_begin)
    p._p.append(r1)

    r2 = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" SEQ {seq_name} \\* ARABIC"
    r2.append(instr)
    p._p.append(r2)

    r3 = OxmlElement("w:r")
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    r3.append(fld_sep)
    p._p.append(r3)

    r4 = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r4.append(t)
    p._p.append(r4)

    r5 = OxmlElement("w:r")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r5.append(fld_end)
    p._p.append(r5)


# ── Outline level ──────────────────────────────────────────────────────


def set_outline_level(p, level: int) -> None:
    """Set the outline level of a heading paragraph."""
    pPr = p._p.get_or_add_pPr()
    for existing in pPr.findall(qn("w:outlineLvl")):
        pPr.remove(existing)
    el = OxmlElement("w:outlineLvl")
    el.set(qn("w:val"), str(level - 1))
    pPr.append(el)


# ── Borders ────────────────────────────────────────────────────────────


def add_bottom_border(p, color: str = "CCCCCC", sz: str = "6") -> None:
    """Add a bottom border to a paragraph."""
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── Page break ─────────────────────────────────────────────────────────


def add_page_break(doc: Document) -> None:
    """Add a page break paragraph."""
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


# ── Table of Contents ──────────────────────────────────────────────────


def add_toc(doc: Document, styles: dict, depth: str = "1-3") -> None:
    """Insert a TOC field at the beginning of the document."""
    toc_fmt = styles.get("toc_title", styles.get("h1", ParagraphFormat()))
    p_title = doc.add_paragraph()
    toc_fmt.apply_to_paragraph(p_title)
    r = p_title.add_run("目录")
    toc_fmt.apply_to_run(r)

    p = doc.add_paragraph()

    def _mk_r(*children) -> OxmlElement:
        r = OxmlElement("w:r")
        for c in children:
            r.append(c)
        return r

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    p._p.append(_mk_r(begin))

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f' TOC \\o "{depth}" \\z \\u '
    p._p.append(_mk_r(instr))

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    p._p.append(_mk_r(separate))

    p_ph = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "（右键此处 → 更新域）"
    p_ph.append(t)
    p._p.append(p_ph)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    p._p.append(_mk_r(end))


# ── Image embedder (SVG + raster) ────────────────────────────────────────


def embed_svg_in_paragraph(
    doc: Document, p, img_data: bytes, alt: str = "",
    max_width_inches: float = 5.5,
) -> None:
    """Embed SVG or raster image in *p*. Auto-detects format from bytes.

    For SVG: embeds via OPC-level ``asvg:svgBlip``.
    For raster images (JPEG/PNG): embeds as standard image blip.
    """
    is_svg = _is_svg_data(img_data)

    if is_svg:
        size = get_svg_size(img_data)
    else:
        from PIL import Image as PILImage
        from io import BytesIO
        try:
            pil_img = PILImage.open(BytesIO(img_data))
            size = pil_img.size  # (width_px, height_px)
        except Exception:
            size = None

    if size is None:
        from .context import ParagraphFormat
        pf = ParagraphFormat()
        _push_run_simple(p, f"[{alt}]", pf, italic=True)
        return

    w_px, h_px = size

    max_px = max_width_inches * 96
    if w_px > max_px:
        scale = max_px / w_px
        w_px = int(w_px * scale)
        h_px = int(h_px * scale)

    emu_w = w_px * 914400 // 96
    emu_h = h_px * 914400 // 96

    if is_svg:
        partname = doc.part.package.next_partname("/word/media/image%d.svg")
        mime = "image/svg+xml"
    elif img_data[:4] == b'\x89PNG':
        partname = doc.part.package.next_partname("/word/media/image%d.png")
        mime = "image/png"
    elif img_data[:2] == b'\xff\xd8':
        partname = doc.part.package.next_partname("/word/media/image%d.jpg")
        mime = "image/jpeg"
    else:
        partname = doc.part.package.next_partname("/word/media/image%d.png")
        mime = "image/png"

    img_part = ImagePart(partname, mime, img_data)
    img_part._package = doc.part.package

    rId = doc.part.relate_to(
        img_part,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image",
    )

    if is_svg:
        xml = _drawing_xml_svg(emu_w, emu_h, rId, alt)
    else:
        xml = _drawing_xml_raster(emu_w, emu_h, rId, alt)

    run = p.add_run()
    run._r.append(parse_xml(xml))


def _is_svg_data(data: bytes) -> bool:
    """Detect SVG data by checking for SVG/XML opening tag."""
    return data[:4] == b'<svg' or data[:5] == b'<?xml'


def _drawing_xml_svg(emu_w: int, emu_h: int, rId: str, alt: str) -> str:
    """Build OOXML drawing element for SVG with asvg:svgBlip extension."""
    return (
        f'<w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        f'  xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
        f'  xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        f'  xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'
        f'  xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"'
        f'  xmlns:asvg="http://schemas.microsoft.com/office/drawing/2020/SVG">'
        f'  <wp:inline distT="0" distB="0" distL="0" distR="0">'
        f'    <wp:extent cx="{emu_w}" cy="{emu_h}"/>'
        f'    <wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'    <wp:docPr id="0" name="SVG" descr="{alt}"/>'
        f'    <wp:cNvGraphicFramePr><a:graphicFrameLocks noChangeAspect="1"/></wp:cNvGraphicFramePr>'
        f'    <a:graphic>'
        f'      <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        f'        <pic:pic>'
        f'          <pic:nvPicPr>'
        f'            <pic:cNvPr id="0" name="SVG"/>'
        f'            <pic:cNvPicPr/>'
        f'          </pic:nvPicPr>'
        f'          <pic:blipFill>'
        f'            <a:blip r:embed="{rId}">'
        f'              <a:extLst>'
        f'                <a:ext uri="{{96DAC541-7B7A-43D3-8B79-5D63384654D6}}">'
        f'                  <asvg:svgBlip r:embed="{rId}"/>'
        f'                </a:ext>'
        f'              </a:extLst>'
        f'            </a:blip>'
        f'            <a:srcRect/>'
        f'            <a:stretch><a:fillRect/></a:stretch>'
        f'          </pic:blipFill>'
        f'          <pic:spPr>'
        f'            <a:xfrm><a:off x="0" y="0"/><a:ext cx="{emu_w}" cy="{emu_h}"/></a:xfrm>'
        f'            <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
        f'          </pic:spPr>'
        f'        </pic:pic>'
        f'      </a:graphicData>'
        f'    </a:graphic>'
        f'  </wp:inline>'
        f'</w:drawing>'
    )


def _drawing_xml_raster(emu_w: int, emu_h: int, rId: str, alt: str) -> str:
    """Build OOXML drawing element for a raster image (JPEG/PNG)."""
    return (
        f'<w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        f'  xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
        f'  xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        f'  xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'
        f'  xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        f'  <wp:inline distT="0" distB="0" distL="0" distR="0">'
        f'    <wp:extent cx="{emu_w}" cy="{emu_h}"/>'
        f'    <wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'    <wp:docPr id="0" name="Picture" descr="{alt}"/>'
        f'    <wp:cNvGraphicFramePr><a:graphicFrameLocks noChangeAspect="1"/></wp:cNvGraphicFramePr>'
        f'    <a:graphic>'
        f'      <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        f'        <pic:pic>'
        f'          <pic:nvPicPr>'
        f'            <pic:cNvPr id="0" name="Picture"/>'
        f'            <pic:cNvPicPr/>'
        f'          </pic:nvPicPr>'
        f'          <pic:blipFill>'
        f'            <a:blip r:embed="{rId}"/>'
        f'            <a:srcRect/>'
        f'            <a:stretch><a:fillRect/></a:stretch>'
        f'          </pic:blipFill>'
        f'          <pic:spPr>'
        f'            <a:xfrm><a:off x="0" y="0"/><a:ext cx="{emu_w}" cy="{emu_h}"/></a:xfrm>'
        f'            <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
        f'          </pic:spPr>'
        f'        </pic:pic>'
        f'      </a:graphicData>'
        f'    </a:graphic>'
        f'  </wp:inline>'
        f'</w:drawing>'
    )


def _push_run_simple(p, text: str, fmt, **overrides) -> None:
    """Minimal push_run for internal use (avoids circular import)."""
    if not text:
        return
    run = p.add_run(text)
    fmt.apply_to_run(run)
    for k, v in overrides.items():
        if k == "italic":
            run.font.italic = v
        elif k == "bold":
            run.font.bold = v


# ── OMML helpers ──────────────────────────────────────────────────────


def insert_inline_omml(p, omml_xml: str) -> None:
    """Insert an inline OMML formula inside a run in *p*."""
    run = p.add_run()
    run._r.append(parse_xml(omml_xml))


def insert_block_omml(doc: Document, omml_xml: str) -> None:
    """Insert a display OMML formula as a centered paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.alignment = 1  # center
    p._p.append(parse_xml(omml_xml))


# ── Table styles ──────────────────────────────────────────────────────


def style_table(table, three_line: bool = False) -> None:
    """Apply border styling to a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    for existing in list(tblPr.findall(qn("w:tblBorders"))):
        tblPr.remove(existing)
    borders = OxmlElement("w:tblBorders")

    if three_line:
        for edge, sz, color in [
            ("top", "12", "000000"),
            ("bottom", "12", "000000"),
            ("insideH", "6", "000000"),
        ]:
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), sz)
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
            borders.append(el)
        for edge in ("left", "right", "insideV"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "none")
            borders.append(el)
    else:
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "999999")
            borders.append(el)
    tblPr.append(borders)


# ── List numbering helpers ────────────────────────────────────────────


_NUM_CACHE_KEY = "_md2word_list_abstracts"


def ensure_list_abstract_nums(doc: Document) -> dict[str, int]:
    """Create abstractNum definitions for bullet and ordered lists if missing.

    Returns dict ``{"bullet": id, "ordered": id}``.
    """
    if hasattr(doc, _NUM_CACHE_KEY):
        return getattr(doc, _NUM_CACHE_KEY)

    numbering_part = doc.part.numbering_part
    numbering_xml = numbering_part._element

    existing_abstract = numbering_xml.findall(qn("w:abstractNum"))
    ids = [
        int(a.get(qn("w:abstractNumId")))
        for a in existing_abstract
        if a.get(qn("w:abstractNumId")) is not None
    ]
    next_id = max(ids) + 1 if ids else 0

    bullet_abstract = (
        f'<w:abstractNum w:abstractNumId="{next_id}"'
        f' xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'  <w:multiLevelType w:val="hybridMultilevel"/>'
        f'  <w:lvl w:ilvl="0"><w:start w:val="1"/><w:numFmt w:val="bullet"/>'
        f'    <w:lvlText w:val="●"/><w:lvlJc w:val="left"/>'
        f'    <w:pPr><w:ind w:left="720" w:hanging="360"/></w:pPr>'
        f'  </w:lvl>'
        f'  <w:lvl w:ilvl="1"><w:start w:val="1"/><w:numFmt w:val="bullet"/>'
        f'    <w:lvlText w:val="○"/><w:lvlJc w:val="left"/>'
        f'    <w:pPr><w:ind w:left="1440" w:hanging="360"/></w:pPr>'
        f'  </w:lvl>'
        f'  <w:lvl w:ilvl="2"><w:start w:val="1"/><w:numFmt w:val="bullet"/>'
        f'    <w:lvlText w:val="▪"/><w:lvlJc w:val="left"/>'
        f'    <w:pPr><w:ind w:left="2160" w:hanging="360"/></w:pPr>'
        f'  </w:lvl>'
        f'  <w:lvl w:ilvl="3"><w:start w:val="1"/><w:numFmt w:val="bullet"/>'
        f'    <w:lvlText w:val="◆"/><w:lvlJc w:val="left"/>'
        f'    <w:pPr><w:ind w:left="2880" w:hanging="360"/></w:pPr>'
        f'  </w:lvl>'
        f'</w:abstractNum>'
    )
    numbering_xml.append(parse_xml(bullet_abstract))
    bullet_id = next_id
    next_id += 1

    ordered_abstract = (
        f'<w:abstractNum w:abstractNumId="{next_id}"'
        f' xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'  <w:multiLevelType w:val="hybridMultilevel"/>'
        f'  <w:lvl w:ilvl="0"><w:start w:val="1"/><w:numFmt w:val="decimal"/>'
        f'    <w:lvlText w:val="%1."/><w:lvlJc w:val="left"/>'
        f'    <w:pPr><w:ind w:left="720" w:hanging="360"/></w:pPr>'
        f'  </w:lvl>'
        f'  <w:lvl w:ilvl="1"><w:start w:val="1"/><w:numFmt w:val="lowerLetter"/>'
        f'    <w:lvlText w:val="%2."/><w:lvlJc w:val="left"/>'
        f'    <w:pPr><w:ind w:left="1440" w:hanging="360"/></w:pPr>'
        f'  </w:lvl>'
        f'  <w:lvl w:ilvl="2"><w:start w:val="1"/><w:numFmt w:val="lowerRoman"/>'
        f'    <w:lvlText w:val="%3."/><w:lvlJc w:val="left"/>'
        f'    <w:pPr><w:ind w:left="2160" w:hanging="360"/></w:pPr>'
        f'  </w:lvl>'
        f'  <w:lvl w:ilvl="3"><w:start w:val="1"/><w:numFmt w:val="decimal"/>'
        f'    <w:lvlText w:val="%4."/><w:lvlJc w:val="left"/>'
        f'    <w:pPr><w:ind w:left="2880" w:hanging="360"/></w:pPr>'
        f'  </w:lvl>'
        f'</w:abstractNum>'
    )
    numbering_xml.append(parse_xml(ordered_abstract))
    ordered_id = next_id

    result = {"bullet": bullet_id, "ordered": ordered_id}
    setattr(doc, _NUM_CACHE_KEY, result)
    return result


def create_list_num_id(doc: Document, list_type: str = "bullet") -> int:
    """Create a new ``<w:num>`` instance and return its numId."""
    abstracts = ensure_list_abstract_nums(doc)
    abstract_id = abstracts[list_type]

    numbering_xml = doc.part.numbering_part._element
    existing_nums = numbering_xml.findall(qn("w:num"))
    num_ids = [
        int(n.get(qn("w:numId")))
        for n in existing_nums
        if n.get(qn("w:numId")) is not None
    ]
    next_num_id = max(num_ids) + 1 if num_ids else 1

    if list_type == "ordered":
        # Each <ol> block must start at 1.  Without an explicit
        # startOverride Word may continue numbering from a previous
        # <w:num> instance that shares the same abstractNum.
        num_xml = (
            f'<w:num xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            f' w:numId="{next_num_id}">'
            f'  <w:abstractNumId w:val="{abstract_id}"/>'
            f'  <w:lvlOverride w:ilvl="0">'
            f'    <w:startOverride w:val="1"/>'
            f'  </w:lvlOverride>'
            f'</w:num>'
        )
    else:
        num_xml = (
            f'<w:num xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            f' w:numId="{next_num_id}">'
            f'  <w:abstractNumId w:val="{abstract_id}"/>'
            f'</w:num>'
        )
    numbering_xml.append(parse_xml(num_xml))
    return next_num_id


def add_num_pr(p, num_id: int, ilvl: int = 0) -> None:
    """Add ``<w:numPr>`` to a paragraph, replacing any existing one."""
    pPr = p._p.get_or_add_pPr()
    for existing in pPr.findall(qn("w:numPr")):
        pPr.remove(existing)
    numPr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    numPr.append(ilvl_el)
    numPr.append(num_id_el)
    pPr.insert(0, numPr)
