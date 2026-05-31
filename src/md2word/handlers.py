"""Block handlers — converts ElementTree blocks into Word paragraphs.

Each ``_handle_*`` function takes a parsed HTML element and manipulates
the python-docx Document to produce the corresponding Word output.
"""

from __future__ import annotations

import base64
import re
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .context import ConversionContext
from .footnotes import create_footnote_reference_run
from .image_utils import (
    get_image_dimensions,
    get_svg_size,
    is_svg_source,
    resolve_image,
    resolve_svg_raw,
)
from .math_omml import latex_to_omml as _latex_to_omml
from .math_renderer import extract_and_placeholder, has_math, render_math_svg
from .mermaid_renderer import render_mermaid
from .ooxml_helpers import (
    add_bookmark,
    add_bottom_border,
    add_num_pr,
    add_page_break,
    add_seq_number,
    add_toc,
    create_list_num_id,
    embed_svg_in_paragraph,
    insert_block_omml,
    insert_inline_omml,
    make_ref_field,
    set_outline_level,
    style_table,
)
from .syntax import extract_language, highlight
from .template import ParagraphFormat, _STYLE_KEYWORDS, set_run_font


# ── List marker regexes (used by ensure_list_blank_lines) ─────────────────

_LIST_MARKER_RE = re.compile(r"^[\-\*\+] ")
_NUMBERED_MARKER_RE = re.compile(r"^\d+[\.\)] ")


# ── HTML pre-processing ────────────────────────────────────────────────


def prepare_html(html: str) -> str:
    """Convert bare \\n in text content to <br/>, preserving <pre> blocks."""
    pres: list[str] = []
    def _save(m):
        pres.append(m.group(0))
        return f"\x00PRE{len(pres)-1}\x00"
    html = re.sub(r'<pre[^>]*>.*?</pre>', _save, html, flags=re.DOTALL)

    html = html.replace('\n', '<br/>')
    html = re.sub(r'>\s*<br/>\s*<', '><', html)
    html = re.sub(r'^(<br/>\s*)+', '', html)
    html = re.sub(r'(<br/>\s*)+$', '', html)

    for i, block in enumerate(pres):
        html = html.replace(f"\x00PRE{i}\x00", block)
    return html


def html_to_blocks(html: str) -> list[ET.Element]:
    root = ET.fromstring(f"<root>{html}</root>")
    return list(root)


# ── Helper: inline text extraction ─────────────────────────────────────


def inline_text(elem: ET.Element) -> str:
    parts: list[str] = []
    if elem.text:
        parts.append(elem.text)
    for child in elem:
        parts.append(inline_text(child))
        if child.tail:
            parts.append(child.tail)
    return "".join(parts)


# ── Helper: Markdown pre-processing for extended syntax ─────────────────


_EXTRA_SYNTAX_PREPROC = [
    # Strikethrough: ~~text~~ → <del>text</del>
    (re.compile(r"(?<!\~)\~\~(.+?)\~\~(?!\~)"), r"<del>\1</del>"),
    # Highlight: ==text== → <mark>text</mark>
    (re.compile(r"\=\=(.+?)\=\="), r"<mark>\1</mark>"),
    # Superscript: ^text^ → <sup>text</sup>
    (re.compile(r"\^(.+?)\^"), r"<sup>\1</sup>"),
    # Subscript: ~text~ → <sub>text</sub> (avoids matching ~~)
    (re.compile(r"(?<!\~)\~(.+?)\~(?!\~)"), r"<sub>\1</sub>"),
]


def preprocess_extended_syntax(text: str) -> str:
    """Pre-process markdown to convert extended syntax to HTML tags.

    Handles: ~~strikethrough~~, ==highlight==, ^sup^, ~sub~
    """
    for pattern, replacement in _EXTRA_SYNTAX_PREPROC:
        text = pattern.sub(replacement, text)
    return text


# ── List blank line insertion ────────────────────────────────────────────


def ensure_list_blank_lines(text: str) -> str:
    """Insert blank lines before list items when missing."""
    lines = text.split("\n")
    result: list[str] = []
    for i, line in enumerate(lines):
        stripped = line.strip()
        is_list = bool(
            _LIST_MARKER_RE.match(line) or _NUMBERED_MARKER_RE.match(line)
        )
        if is_list and i > 0:
            prev = lines[i - 1].strip()
            if (
                prev
                and not _LIST_MARKER_RE.match(lines[i - 1])
                and not _NUMBERED_MARKER_RE.match(lines[i - 1])
                and not prev.startswith("#")
            ):
                if result and result[-1] != "":
                    result.append("")
        result.append(line)
    return "\n".join(result)


# ── Math post-processing ───────────────────────────────────────────────


def postprocess_math_html(html: str, math_exprs: list[dict]) -> str:
    """Render math LaTeX to OMML (preferred) or SVG, replacing placeholders."""
    result = html
    for expr in math_exprs:
        latex = expr["latex"]
        kind = expr["kind"]
        from xml.sax.saxutils import escape as _xml_escape
        alt = _xml_escape(f"[{latex[:60]}]")

        omml = _latex_to_omml(latex, kind == "block")
        if omml:
            omml_b64 = base64.b64encode(omml.encode("utf-8")).decode("ascii")
            replacement = (
                f'<img class="math-{kind}" src="" '
                f'data-omml-base64="{omml_b64}" alt="{alt}"/>'
            )
        else:
            svg_bytes, w_px, h_px = render_math_svg(latex, kind)
            if svg_bytes:
                svg_b64 = base64.b64encode(svg_bytes).decode("ascii")
                replacement = (
                    f'<img class="math-{kind}" src="" '
                    f'data-svg-base64="{svg_b64}" alt="{alt}"/>'
                )
            else:
                replacement = f'<span class="math-fallback">{alt}</span>'
        result = result.replace(expr["placeholder"], replacement)
    return result


# ── Run helpers ────────────────────────────────────────────────────────


def _set_run_color(run, color: RGBColor) -> None:
    run.font.color.rgb = color


def push_run(p, text: str, base_fmt: ParagraphFormat, **overrides) -> None:
    if not text:
        return
    run = p.add_run(text)
    base_fmt.apply_to_run(run)
    for k, v in overrides.items():
        if k == "color":
            _set_run_color(run, v)
        elif k == "name":
            set_run_font(run, v, v)
        elif k == "underline":
            run.font.underline = v
        elif k == "strikethrough":
            run.font.strike = v
        elif k == "highlight_color":
            run.font.highlight_color = v
        elif k == "superscript":
            run.font.superscript = v
        elif k == "subscript":
            run.font.subscript = v
        else:
            setattr(run.font, k, v)


def push_text_with_footnotes(
    p, text: str, base_fmt: ParagraphFormat,
    fn_map: dict[str, int],
    **overrides,
) -> None:
    """Push *text*, splitting on footnote placeholders."""
    from .converter import _FN_PLACEHOLDER_RE
    parts = _FN_PLACEHOLDER_RE.split(text)
    for i, part in enumerate(parts):
        if not part:
            continue
        if i % 2 == 0:
            push_run(p, part, base_fmt, **overrides)
        else:
            fn_id = fn_map.get(part)
            if fn_id is not None:
                p._p.append(create_footnote_reference_run(None, fn_id))
            else:
                push_run(p, f"[{part}]", base_fmt, **overrides)


def _detect_task_prefix(text: str) -> tuple[str, str | None]:
    """Detect task list checkbox prefix.

    Returns ``(remaining_text, checkbox_char)`` where *checkbox_char* is
    ``"☑ "`` for checked, ``"☐ "`` for unchecked, or ``None`` if no marker.
    """
    if text.startswith("[x] ") or text.startswith("[X] "):
        return text[4:], "☑ "
    elif text.startswith("[ ] "):
        return text[4:], "☐ "
    return text, None


def _push_detect(p, text: str, base_fmt: ParagraphFormat, **overrides):
    """Push text, detecting checkboxes and task list markers."""
    if not text:
        return
    remaining, checkbox = _detect_task_prefix(text)
    if checkbox:
        push_run(p, checkbox, base_fmt, color=RGBColor(0x2E, 0x7D, 0x32)
                 if checkbox == "☑ " else RGBColor(0x99, 0x99, 0x99))
        text = remaining
    push_run(p, text, base_fmt, **overrides)


def _make_push_detect(p, base_fmt):
    """Return a closure-based _push_detect that captures *p* and *base_fmt*."""
    def _fn(text: str, **overrides):
        if not text:
            return
        remaining, checkbox = _detect_task_prefix(text)
        if checkbox:
            push_run(p, checkbox, base_fmt, color=RGBColor(0x2E, 0x7D, 0x32)
                     if checkbox == "☑ " else RGBColor(0x99, 0x99, 0x99))
            text = remaining
        push_run(p, text, base_fmt, **overrides)
    return _fn


# ── Build runs ─────────────────────────────────────────────────────────


def _build_inline_runs_core(
    doc: Document, p, elem: ET.Element,
    base_fmt: ParagraphFormat,
    fn_map: dict[str, int] | None = None,
    skip_tags: set[str] | None = None,
    enable_footnotes: bool = True,
    enable_cross_ref: bool = True,
) -> None:
    """Core inline run builder — handles strong, em, code, a, fn, img, del, mark, sup, sub.

    *skip_tags*: child tags to skip (only tail text captured — for nested lists).
    *enable_footnotes*: process ``<fn>`` tags.
    *enable_cross_ref*: ``[text](#anchor)`` → REF field (else plain hyperlink).
    """
    _pd = _make_push_detect(p, base_fmt)

    if elem.text:
        _pd(elem.text)

    for child in elem:
        tag = child.tag
        if tag == "br":
            continue

        if skip_tags and tag in skip_tags:
            if child.tail:
                _pd(child.tail)
            continue

        if tag in ("strong", "b"):
            _pd(child.text or "", bold=True)
        elif tag in ("em", "i"):
            _pd(child.text or "", italic=True)
        elif tag == "code":
            _pd(child.text or "", name="DengXian")
        elif tag == "a":
            href = child.get("href", "")
            txt = child.text or href
            if enable_cross_ref and href.startswith("#"):
                target = href[1:]
                for ref_run in make_ref_field(txt, target):
                    p._p.append(ref_run)
            else:
                _pd(txt, underline=True, color=RGBColor(0x00, 0x52, 0xCC))
        elif tag == "fn" and enable_footnotes and fn_map:
            fn_id_attr = child.get("id", "")
            if fn_id_attr in fn_map:
                p._p.append(create_footnote_reference_run(None, fn_map[fn_id_attr]))
        elif tag == "img":
            _build_runs_img_inline(doc, p, child, base_fmt)
        elif tag == "del":
            _pd(child.text or "", strikethrough=True)
        elif tag == "mark":
            _pd(child.text or "", highlight_color=7)
        elif tag == "sup":
            _pd(child.text or "", superscript=True)
        elif tag == "sub":
            _pd(child.text or "", subscript=True)
        else:
            _pd(child.text or "")

        if child.tail:
            _pd(child.tail)


def build_runs(
    doc: Document,
    p,
    elem: ET.Element,
    base_fmt: ParagraphFormat,
    fn_map: dict[str, int] | None = None,
) -> None:
    """Populate a docx paragraph with runs reflecting inline HTML tags."""
    _build_inline_runs_core(doc, p, elem, base_fmt, fn_map=fn_map,
                            enable_footnotes=True, enable_cross_ref=True)


def _build_runs_img_inline(doc, p, child, base_fmt):
    """Handle an <img> inside a paragraph (inline math/image)."""
    src = child.get("src", "")
    alt = child.get("alt", "")
    omml_b64 = child.get("data-omml-base64", "")
    if omml_b64:
        omml_data = base64.b64decode(omml_b64).decode("utf-8")
        if omml_data:
            insert_inline_omml(p, omml_data)
        else:
            push_run(p, f"[Math: {alt}]", base_fmt, italic=True)
    elif svg_b64 := child.get("data-svg-base64", ""):
        svg_data = base64.b64decode(svg_b64)
        if svg_data:
            embed_svg_in_paragraph(doc, p, svg_data, alt, 3.5)
        else:
            push_run(p, f"[Math: {alt}]", base_fmt, italic=True)
    elif is_svg_source(src):
        svg_data = resolve_svg_raw(src)
        if svg_data:
            embed_svg_in_paragraph(doc, p, svg_data, alt, 4.0)
        else:
            push_run(p, f"[SVG: {alt}]", base_fmt, italic=True)
    else:
        stream = resolve_image(src)
        if stream is not None:
            w, h = get_image_dimensions(stream, max_width_inches=4.0)
            run = p.add_run()
            run.add_picture(stream, width=w, height=h)
        else:
            push_run(p, f"[Image: {alt}]", base_fmt, italic=True)


def split_by_br(elem: ET.Element) -> list[tuple[str, ET.Element | None]]:
    """Split element content by <br/> tags.

    Returns list of (text, img_element) pairs — *img_element* is ``None``
    for plain-text segments and the ``<img>`` element for image-only segments.
    """
    segments: list[tuple[str, ET.Element | None]] = []
    buf: list[str] = []

    def _flush():
        nonlocal buf
        text = "".join(buf).strip()
        if text:
            segments.append((text, None))
        buf = []

    if elem.text:
        buf.append(elem.text)
    for child in elem:
        if child.tag == "br":
            _flush()
        elif child.tag == "img":
            _flush()
            segments.append(("", child))
        else:
            buf.append(inline_text(child))
        if child.tail:
            buf.append(child.tail)
    _flush()
    if not segments:
        segments = [("", None)]
    return segments


def build_runs_skip(doc, p, elem, base_fmt, skip_tags=None):
    """Like build_runs but skips child tags in *skip_tags* (for nested lists)."""
    _build_inline_runs_core(doc, p, elem, base_fmt,
                            skip_tags=set(skip_tags) if skip_tags else None,
                            enable_footnotes=False, enable_cross_ref=False)


# ── Footnote placeholder ───────────────────────────────────────────────


def replace_fn_placeholders(html: str) -> str:
    """Replace \\x00FN_id\\x00 → XML-safe <fn id="id"/> tags."""
    from .converter import _FN_PLACEHOLDER_RE, _FN_TAG_RE
    return _FN_PLACEHOLDER_RE.sub(r'<fn id="\1"/>', html)


# ── Token helpers for code highlighting ────────────────────────────────


def tokens_for_line(tokens: list, line_text: str) -> list:
    """Filter tokens to match a single line."""
    line = line_text.rstrip("\n")
    result = []
    offset = 0
    for tok in tokens:
        if offset >= len(line):
            break
        idx = line.find(tok.text, offset)
        if idx == -1:
            continue
        if idx > offset:
            result.append(type(tok)(text=line[offset:idx], color=None, bold=False, italic=False))
        result.append(tok)
        offset = idx + len(tok.text)
    if offset < len(line):
        result.append(type(tok)(text=line[offset:], color=None, bold=False, italic=False))
    return result


# ── Slot slug for bookmark ─────────────────────────────────────────────


def _slugify(text: str) -> str:
    """Convert text to a bookmark-friendly ASCII slug."""
    text = text.strip().lower()
    text = re.sub(r"[^\w\s-]", "", text)
    text = re.sub(r"[\s_]+", "-", text)
    text = re.sub(r"-+", "-", text)
    return text.strip("-") or "ref"


# ── Block handlers ─────────────────────────────────────────────────────


def handle_heading(
    doc: Document, elem: ET.Element, level: int, styles: dict,
    ctx: ConversionContext,
    number_headings: bool = False,
    page_break: bool = False,
) -> None:
    slot = styles.get(f"h{level}") or styles.get("body", ParagraphFormat())
    if page_break and level == 1:
        add_page_break(doc)
    p = doc.add_paragraph()
    # Apply Word heading style so TOC \o can resolve it
    style_name = f"Heading {level}"
    try:
        p.style = doc.styles[style_name]
    except KeyError:
        pass
    slot.apply_to_paragraph(p)
    if number_headings:
        num = ctx.next_heading_number(level)
        if num:
            push_run(p, f"{num} ", slot)
    build_runs(doc, p, elem, slot, fn_map=ctx.fn_map if ctx.fn_map else None)
    set_outline_level(p, level)

    # Add bookmark from heading text for cross-references
    heading_text = inline_text(elem)
    if heading_text:
        slug = _slugify(heading_text)
        slug = ctx.unique_bookmark_slug(slug)
        add_bookmark(p, ctx, name=slug)


def handle_paragraph(
    doc: Document, elem: ET.Element, styles: dict,
    ctx: ConversionContext,
    image_max_width: float = 5.5,
) -> None:
    """Normal paragraph — detects lone images and embedded <br>."""
    imgs = elem.findall("img")
    if len(imgs) == 1 and not elem.text and not any(
        c.tag != "img" for c in elem
    ):
        return handle_image(doc, imgs[0], styles, ctx, image_max_width=image_max_width)

    fmt = styles.get("body", ParagraphFormat())

    brs = elem.findall(".//br")
    if brs:
        for text_part, img_part in split_by_br(elem):
            if img_part is not None:
                handle_image(doc, img_part, styles, ctx, image_max_width=image_max_width)
            elif text_part:
                p = doc.add_paragraph()
                fmt.apply_to_paragraph(p)
                push_run(p, text_part, fmt)
        return

    p = doc.add_paragraph()
    fmt.apply_to_paragraph(p)
    build_runs(doc, p, elem, fmt, fn_map=ctx.fn_map if ctx.fn_map else None)


def handle_image(
    doc: Document, elem: ET.Element, styles: dict,
    ctx: ConversionContext,
    image_max_width: float = 5.5,
) -> None:
    """Image with alt-text caption below. Supports raster and SVG."""
    src = elem.get("src", "")
    alt = elem.get("alt", "")

    omml_b64 = elem.get("data-omml-base64", "")
    if omml_b64:
        omml_data = base64.b64decode(omml_b64).decode("utf-8")
        if omml_data:
            insert_block_omml(doc, omml_data)
        else:
            p = doc.add_paragraph()
            push_run(p, f"[Math: {alt}]", ParagraphFormat(), italic=True)
        return

    img_fmt = styles.get("image", ParagraphFormat())
    is_math = bool(elem.get("data-svg-base64", ""))

    svg_b64 = elem.get("data-svg-base64", "")
    if svg_b64:
        svg_data = base64.b64decode(svg_b64)
        if not svg_data:
            p = doc.add_paragraph()
            push_run(p, f"[Math: {alt}]", ParagraphFormat(), italic=True)
            return
        p = doc.add_paragraph()
        if not is_math:
            img_fmt.apply_to_paragraph(p)
        embed_svg_in_paragraph(doc, p, svg_data, alt, max_width_inches=image_max_width)
    elif is_svg_source(src):
        svg_data = resolve_svg_raw(src)
        if svg_data is None:
            p = doc.add_paragraph()
            push_run(p, f"[SVG: {alt}]", ParagraphFormat(), italic=True)
            return
        p = doc.add_paragraph()
        img_fmt.apply_to_paragraph(p)
        embed_svg_in_paragraph(doc, p, svg_data, alt, max_width_inches=image_max_width)
    else:
        stream = resolve_image(src)
        if stream is None:
            p = doc.add_paragraph()
            push_run(p, f"[Image: {alt}]", ParagraphFormat(), italic=True)
            return
        p = doc.add_paragraph()
        img_fmt.apply_to_paragraph(p)
        width, height = get_image_dimensions(stream, max_width_inches=image_max_width)
        run = p.add_run()
        run.add_picture(stream, width=width, height=height)

    if alt and not is_math:
        cap_fmt = styles.get("figcaption", ParagraphFormat())
        if cap_fmt.font_name is None:
            p_cap = doc.add_paragraph()
            p_cap.paragraph_format.alignment = 1
            r = p_cap.add_run(alt)
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            set_run_font(r, "SimSun", "SimSun")
        else:
            p_cap = doc.add_paragraph()
            cap_fmt.apply_to_paragraph(p_cap)
            r = p_cap.add_run(alt)
            cap_fmt.apply_to_run(r)
        add_bookmark(p_cap, ctx, f"fig-{_slugify(alt)}")


def handle_blockquote(
    doc: Document, elem: ET.Element, styles: dict,
    ctx: ConversionContext,
) -> None:
    """Blockquote — process child <p> elements."""
    fmt = styles.get("quote", styles.get("body", ParagraphFormat()))
    for child in elem:
        if child.tag == "p":
            p = doc.add_paragraph()
            fmt.apply_to_paragraph(p)
            build_runs(doc, p, child, fmt, fn_map=ctx.fn_map if ctx.fn_map else None)
        else:
            p = doc.add_paragraph()
            fmt.apply_to_paragraph(p)
            build_runs(doc, p, child if child.tag else elem, fmt, fn_map=ctx.fn_map if ctx.fn_map else None)


def handle_code_block(
    doc: Document, elem: ET.Element, styles: dict,
    ctx: ConversionContext,
    highlight_enabled: bool = True,
    mermaid_enabled: bool = True,
) -> None:
    """Code block — each line is a paragraph, with optional syntax highlighting."""
    fmt = styles.get("code", ParagraphFormat())
    code_elem = elem.find("code")
    lang = extract_language(code_elem) if code_elem is not None else ""

    text = inline_text(elem)

    if mermaid_enabled and lang == "mermaid" and text.strip():
        result = render_mermaid(text.strip())
        if result:
            img_fmt = styles.get("image", ParagraphFormat())
            p = doc.add_paragraph()
            img_fmt.apply_to_paragraph(p)
            embed_svg_in_paragraph(doc, p, result.svg_bytes, "mermaid", max_width_inches=5.5)
            return
        else:
            print("  [WARN] Falling back to plain text for mermaid block")

    tokens = None
    if highlight_enabled and lang != "mermaid":
        tokens = highlight(text, lang)

    for line in text.strip('\n').split('\n'):
        p = doc.add_paragraph()
        fmt.apply_to_paragraph(p)
        if not line:
            run = p.add_run(" ")
            fmt.apply_to_run(run)
            set_run_font(run, fmt.font_name or "DengXian", fmt.east_asia_font or "DengXian")
            continue

        if tokens:
            line_tokens = tokens_for_line(tokens, line)
            for tok in line_tokens:
                run = p.add_run(tok.text)
                fmt.apply_to_run(run)
                set_run_font(run, fmt.font_name or "Consolas", fmt.east_asia_font or "DengXian")
                if tok.color:
                    run.font.color.rgb = tok.color
                if tok.bold:
                    run.font.bold = True
                if tok.italic:
                    run.font.italic = True
        else:
            run = p.add_run(line)
            fmt.apply_to_run(run)
            set_run_font(run, fmt.font_name or "DengXian", fmt.east_asia_font or "DengXian")


def handle_unordered_list(
    doc: Document, elem: ET.Element, styles: dict,
    ctx: ConversionContext,
    depth: int = 0,
    num_id: int | None = None,
) -> None:
    """Unordered list with Word native bullet numbering. Supports nesting."""
    fmt = styles.get("bullet_list", styles.get("body", ParagraphFormat()))
    if num_id is None:
        num_id = create_list_num_id(doc, "bullet")
    for li in elem:
        if li.tag != "li":
            continue
        nested = li.findall("ul") + li.findall("ol")
        p = doc.add_paragraph()
        fmt.apply_to_paragraph(p)
        add_num_pr(p, num_id, ilvl=depth)
        build_runs_skip(doc, p, li, fmt, skip_tags={"ul", "ol"})
        if nested:
            for nest in nested:
                if nest.tag == "ul":
                    handle_unordered_list(doc, nest, styles, ctx, depth + 1, num_id)
                else:
                    handle_ordered_list(doc, nest, styles, ctx, depth + 1)


def handle_ordered_list(
    doc: Document, elem: ET.Element, styles: dict,
    ctx: ConversionContext,
    depth: int = 0,
    num_id: int | None = None,
) -> None:
    """Ordered list with Word native decimal numbering. Supports nesting."""
    fmt = styles.get("number_list", styles.get("body", ParagraphFormat()))
    if num_id is None:
        num_id = create_list_num_id(doc, "ordered")
    for child in elem:
        if child.tag != "li":
            continue
        nested = child.findall("ul") + child.findall("ol")
        p = doc.add_paragraph()
        fmt.apply_to_paragraph(p)
        add_num_pr(p, num_id, ilvl=depth)
        build_runs_skip(doc, p, child, fmt, skip_tags={"ul", "ol"})
        if nested:
            for nest in nested:
                if nest.tag == "ul":
                    handle_unordered_list(doc, nest, styles, ctx, depth + 1)
                else:
                    handle_ordered_list(doc, nest, styles, ctx, depth + 1, num_id)


def handle_table(
    doc: Document, elem: ET.Element, styles: dict,
    ctx: ConversionContext,
    three_line: bool = False,
    table_idx: int = 0,
) -> None:
    """Markdown table."""
    rows_elem = elem.findall(".//tr")
    if not rows_elem:
        return
    cols = max(
        len(tr.findall("th")) + len(tr.findall("td")) for tr in rows_elem
    )
    if cols == 0:
        return

    table = doc.add_table(rows=len(rows_elem), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    body_fmt = styles.get("body", ParagraphFormat())

    for ri, tr in enumerate(rows_elem):
        cells = tr.findall("th") + tr.findall("td")
        for ci, cell in enumerate(cells):
            if ci >= cols:
                break
            wc = table.rows[ri].cells[ci]
            wc.text = ""
            p = wc.paragraphs[0]
            if cell.tag == "th":
                body_fmt.apply_to_paragraph(p)
                push_run(p, inline_text(cell), body_fmt, bold=True)
            else:
                body_fmt.apply_to_paragraph(p)
                build_runs(doc, p, cell, body_fmt, fn_map=ctx.fn_map if ctx.fn_map else None)
    style_table(table, three_line=three_line)
    add_bookmark(table.rows[0].cells[0].paragraphs[0], ctx, f"tbl-{table_idx}")


def handle_horizontal_rule(doc: Document, styles: dict) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    add_bottom_border(p)
