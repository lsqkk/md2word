"""Document metadata — OOXML post-processing, GB compliance, red-head, page numbers."""

from __future__ import annotations

import io
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

from .context import ConversionReport
from .ooxml_helpers import add_bottom_border
from .template import ParagraphFormat, set_run_font


# ── GB Standards reference ─────────────────────────────────────────────────


_GB_STANDARDS: dict[str, dict] = {
    "9704_2012": {
        "margins_cm": (3.7, 3.5, 2.8, 2.6),
        "body_font": "FangSong",
        "body_size_pt": 16,
        "heading1_font": "SimHei",
        "heading1_size_pt": 22,
        "line_spacing": 1.5,
    },
    "7713_2015": {
        "margins_cm": (2.54, 2.54, 3.17, 3.17),
        "body_font": "SimSun",
        "body_size_pt": 12,
        "heading1_font": "SimHei",
        "heading1_size_pt": 16,
        "heading_align": 1,
        "line_spacing": 1.5,
    },
}


# ── OOXML metadata fix ──────────────────────────────────────────────────


def fix_ooxml_metadata(output_path: str | Path) -> None:
    """Post-process docx ZIP to fix thumbnail + application name.

    python-docx embeds a blank ``docProps/thumbnail.jpeg`` that makes
    Windows show a white box instead of a content preview.  We strip
    it so Windows generates a preview from the actual document content.

    Also fixes the Application name from "Microsoft Macintosh Word"
    to "Microsoft Office Word".
    """
    path = Path(output_path)
    buf = path.read_bytes()
    out_buf = io.BytesIO()
    changed = False

    with zipfile.ZipFile(io.BytesIO(buf)) as zin:
        with zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if "thumbnail" in item.filename:
                    changed = True
                    continue

                raw = zin.read(item.filename)
                if item.filename == "docProps/app.xml":
                    text = raw.decode("utf-8")
                    fixed = text.replace(
                        "<Application>Microsoft Macintosh Word</Application>",
                        "<Application>Microsoft Office Word</Application>",
                    )
                    if fixed != text:
                        changed = True
                    raw = fixed.encode("utf-8")
                elif item.filename == "_rels/.rels":
                    without_thumb = re.sub(
                        r'<Relationship[^>]*thumbnail[^>]*/>',
                        '',
                        raw.decode("utf-8"),
                    )
                    if without_thumb != raw.decode("utf-8"):
                        changed = True
                    raw = without_thumb.encode("utf-8")
                zout.writestr(item, raw)

    if changed:
        path.write_bytes(out_buf.getvalue())


# ── GB compliance check ──────────────────────────────────────────────────


def check_gb_compliance(doc: Document, styles: dict, report: ConversionReport) -> None:
    """Check document formatting against Chinese GB standards, emit warnings."""
    if not doc.sections:
        return
    s = doc.sections[0]

    top = s.top_margin / 914400 * 25.4
    bottom = s.bottom_margin / 914400 * 25.4
    left = s.left_margin / 914400 * 25.4
    right = s.right_margin / 914400 * 25.4

    gb = _GB_STANDARDS["9704_2012"]
    gb_t, gb_b, gb_l, gb_r = [v * 10 for v in gb["margins_cm"]]
    checks = [
        (top, gb_t, 2.0, f"上边距 ({top:.0f}mm) 偏离 GB/T 9704-2012 标准 ({gb_t:.0f}mm)"),
        (bottom, gb_b, 2.0, f"下边距 ({bottom:.0f}mm) 偏离 GB/T 9704-2012 标准 ({gb_b:.0f}mm)"),
        (left, gb_l, 2.0, f"左边距 ({left:.0f}mm) 偏离 GB/T 9704-2012 标准 ({gb_l:.0f}mm)"),
        (right, gb_r, 2.0, f"右边距 ({right:.0f}mm) 偏离 GB/T 9704-2012 标准 ({gb_r:.0f}mm)"),
    ]
    for actual, expected, tol, msg in checks:
        if abs(actual - expected) > tol:
            report.info_msg(f"[GB] 非公文模板: {msg}")

    body_fmt = styles.get("body")
    if body_fmt and body_fmt.font_name:
        if body_fmt.font_name not in ("FangSong", "SimSun", "SimHei"):
            report.info_msg(f"[GB] 正文字体 '{body_fmt.font_name}' — 非标准公文/学术字体（推荐仿宋/宋体）")


# ── Guide-paragraph removal ──────────────────────────────────────────


def remove_guide_paragraphs(doc: Document) -> None:
    """Remove all guide paragraphs (style markers) from a template."""
    from .template import _STYLE_KEYWORDS

    to_remove = []
    for p in doc.paragraphs:
        text = p.text.strip().lower()
        if not text:
            continue
        for keyword in sorted(_STYLE_KEYWORDS, key=len, reverse=True):
            if keyword in text:
                to_remove.append(p._p)
                break
    for p_elem in to_remove:
        p_elem.getparent().remove(p_elem)


# ── Red-head document header ────────────────────────────────────────────


def insert_redhead_header(doc: Document, authority_name: str, styles: dict) -> None:
    """Insert 红头文件 header elements at the document start."""
    p_red = doc.add_paragraph()
    p_red.paragraph_format.alignment = 1
    p_red.paragraph_format.space_after = Pt(4)
    run = p_red.add_run(authority_name)
    run.font.name = "SimHei"
    run.font.size = Pt(28)
    run.font.bold = True
    set_run_font(run, "SimHei", "SimHei")
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    p_suffix = doc.add_paragraph()
    p_suffix.paragraph_format.alignment = 1
    p_suffix.paragraph_format.space_after = Pt(6)
    run2 = p_suffix.add_run("文件")
    run2.font.name = "SimHei"
    run2.font.size = Pt(28)
    run2.font.bold = True
    set_run_font(run2, "SimHei", "SimHei")
    run2.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    p_line = doc.add_paragraph()
    p_line.paragraph_format.space_before = Pt(2)
    p_line.paragraph_format.space_after = Pt(12)
    add_bottom_border(p_line, color="CC0000", sz="16")

    p_num = doc.add_paragraph()
    p_num.paragraph_format.alignment = 1
    p_num.paragraph_format.space_after = Pt(12)
    run3 = p_num.add_run("〔2024〕 号")
    run3.font.name = "FangSong"
    run3.font.size = Pt(16)
    set_run_font(run3, "FangSong", "FangSong")


# ── Page number formatting ─────────────────────────────────────────────


def set_page_number_format(doc: Document, fmt: str = "-- %d --") -> None:
    """Set page number in document footer.

    *fmt* uses ``%d`` as placeholder for the page number.
    Example: ``-- %d --`` → ``-- 1 --``
    """
    if not doc.sections:
        return

    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.paragraph_format.alignment = 1

    parts = fmt.split("%d", 1)
    prefix = parts[0]
    suffix = parts[1] if len(parts) > 1 else ""

    if prefix:
        run_pre = p.add_run(prefix)
        run_pre.font.size = Pt(10)

    run_field = OxmlElement("w:r")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run_field.append(fld_begin)
    p._p.append(run_field)

    run_instr = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run_instr.append(instr)
    p._p.append(run_instr)

    run_sep = OxmlElement("w:r")
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run_sep.append(fld_sep)
    p._p.append(run_sep)

    run_disp = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    run_disp.append(t)
    p._p.append(run_disp)

    run_end = OxmlElement("w:r")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_end.append(fld_end)
    p._p.append(run_end)

    if suffix:
        run_suf = p.add_run(suffix)
        run_suf.font.size = Pt(10)
