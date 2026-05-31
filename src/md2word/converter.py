"""Core converter — markdown → Word document using template-derived styles.

This module now serves as the orchestration layer.  Implementation details
have been split into focused sub-modules:

- ``handlers.py``: Block handlers and inline run builders
- ``ooxml_helpers.py``: OOXML-level operations (bookmarks, fields, SVG, etc.)
- ``metadata.py``: Post-processing, GB compliance, red-head, page numbers
- ``context.py``: ConversionContext, ConversionReport with severity levels
- ``frontmatter.py``: YAML front matter parsing and application
- ``cache.py``: Incremental conversion cache
- ``options.py``: ConvertOptions dataclass
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

import markdown
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .cache import file_hash as _file_hash
from .cache import load_cache as _load_cache
from .cache import save_cache as _save_cache
from .config import load_config as _load_config
from .context import ConversionContext, ConversionReport
from .footnotes import add_footnotes_to_document, extract_footnotes
from .frontmatter import apply_front_matter, parse_front_matter
from .handlers import (
    build_runs_skip,
    ensure_list_blank_lines,
    handle_blockquote,
    handle_code_block,
    handle_heading,
    handle_horizontal_rule,
    handle_image,
    handle_ordered_list,
    handle_paragraph,
    handle_table,
    handle_unordered_list,
    html_to_blocks,
    postprocess_math_html,
    preprocess_extended_syntax,
    prepare_html,
    replace_fn_placeholders,
)
from .math_renderer import extract_and_placeholder
from .metadata import (
    check_gb_compliance,
    fix_ooxml_metadata,
    insert_redhead_header,
    remove_guide_paragraphs,
    set_page_number_format,
)
from .options import ConvertOptions
from .template import ParagraphFormat, extract_template_styles


# ── Regex ────────────────────────────────────────────────────────────

_FN_PLACEHOLDER_RE = re.compile(r"\x00FN_([^\x00]+)\x00")
_FN_TAG_RE = re.compile(r'<fn\s+id="([^"]+)"\s*/>')


# ── Helper: insert abstract / keywords from front matter ────────────────


def _insert_abstract_keywords(doc, front_matter, styles):
    """Insert abstract and keywords from YAML front matter into the document."""
    abstract = front_matter.get("abstract")
    keywords = front_matter.get("keywords")

    if abstract:
        label_fmt = styles.get("abstract")
        if label_fmt:
            p = doc.add_paragraph()
            label_fmt.apply_to_paragraph(p)
            r = p.add_run("摘要")
            label_fmt.apply_to_run(r)
        p = doc.add_paragraph()
        fmt = styles.get("body", ParagraphFormat())
        fmt.apply_to_paragraph(p)
        fmt.apply_to_run(p.add_run(abstract))

    if keywords:
        label_fmt = styles.get("keywords")
        if label_fmt:
            p = doc.add_paragraph()
            label_fmt.apply_to_paragraph(p)
            r = p.add_run("关键词")
            label_fmt.apply_to_run(r)
        p = doc.add_paragraph()
        fmt = styles.get("body", ParagraphFormat())
        fmt.apply_to_paragraph(p)
        fmt.apply_to_run(p.add_run(keywords))


# ── Main convert function ────────────────────────────────────────────


def convert(
    markdown_text: str,
    template_path: str | Path,
    output_path: str | Path,
    options: ConvertOptions | None = None,
    **kwargs,
) -> ConversionReport:
    """Convert *markdown_text* to a Word document.

    Args:
        markdown_text: Raw markdown content.
        template_path: Path to the .docx template with guide paragraphs.
        output_path: Where to write the result.
        options: A :class:`ConvertOptions` instance.  If omitted, defaults
                 are used.  For backward compatibility you may also pass
                 individual keyword arguments (e.g. ``toc=False``); these
                 take priority over fields in *options*.
        **kwargs: Legacy keyword arguments — merged into *options* if given.
    """
    # Build ConvertOptions: default → options object → kwargs override
    opts = ConvertOptions()
    if options is not None:
        for fld in opts.__dataclass_fields__:
            val = getattr(options, fld, None)
            if val is not None:
                setattr(opts, fld, val)
    for k, v in kwargs.items():
        if k in opts.__dataclass_fields__ and v is not None:
            setattr(opts, k, v)

    ctx = ConversionContext()
    ctx.report = ConversionReport()
    ctx.styles = extract_template_styles(template_path)
    ctx.styles["_report"] = ctx.report

    # Apply style_map: override slot styles with user-provided mappings
    if opts.style_map:
        for slot, custom_slot in opts.style_map.items():
            if custom_slot in ctx.styles:
                ctx.styles[slot] = ctx.styles[custom_slot]

    # ── Verbose helper ─────────────────────────────────────────────────
    def _vprint(*a, **kw):
        if opts.verbose:
            print("[md2word]", *a, file=sys.stderr, **kw)

    _vprint(f"模板: {template_path}")
    _vprint("解析 YAML front matter…")
    front_matter, markdown_text = parse_front_matter(markdown_text)

    _vprint("预处理扩展语法 (~~strikethrough~~, ==highlight==, etc.)…")
    markdown_text = preprocess_extended_syntax(markdown_text)

    _vprint("预处理脚注…")
    fn_list: list = []
    if opts.footnotes_enabled:
        markdown_text, fn_list = extract_footnotes(markdown_text)

    _vprint("预处理数学公式…")
    if opts.math_enabled:
        markdown_text, math_exprs = extract_and_placeholder(markdown_text)
    else:
        math_exprs = []

    _vprint("确保列表前有空行…")
    markdown_text = ensure_list_blank_lines(markdown_text)

    _vprint(f"Markdown → HTML (扩展: extra, tables, fenced_code, sane_lists)…")
    html = markdown.markdown(
        markdown_text,
        extensions=[
            "extra",
            "tables",
            "fenced_code",
            "sane_lists",
        ],
    )

    html = prepare_html(html)

    _vprint(f"后处理数学公式 (OMML/SVG)…")
    if math_exprs:
        html = postprocess_math_html(html, math_exprs)

    _vprint(f"替换脚注占位符…")
    if fn_list:
        html = replace_fn_placeholders(html)

    blocks = html_to_blocks(html)
    _vprint(f"解析到 {len(blocks)} 个内容块")

    doc = Document(str(template_path))

    _vprint("应用 front matter 到 docx 属性…")
    apply_front_matter(doc, front_matter)

    # Ensure first section doesn't force a new page
    if doc.sections:
        doc.sections[0].start_type = 0

    _vprint("移除引导段落…")
    remove_guide_paragraphs(doc)
    if doc.paragraphs and not doc.paragraphs[0].text.strip():
        p = doc.paragraphs[0]._p
        p.getparent().remove(p)

    _vprint("插入摘要/关键词…")
    _insert_abstract_keywords(doc, front_matter, ctx.styles)

    if opts.redhead_authority:
        _vprint(f"插入红头文件…")
        insert_redhead_header(doc, opts.redhead_authority, ctx.styles,
                               year=opts.redhead_year or 2024,
                               number=opts.redhead_number or "")

    if opts.page_number_fmt:
        _vprint(f"设置页码格式: {opts.page_number_fmt}")
        set_page_number_format(doc, opts.page_number_fmt)

    if opts.gb_check:
        _vprint("检查 GB 标准合规性…")
        check_gb_compliance(doc, ctx.styles, ctx.report)

    from .ooxml_helpers import add_toc as _add_toc

    _toc_pending = opts.toc
    _toc_inserted = False
    check_blocks = list(blocks)
    if _toc_pending and not any(b.tag == "h1" for b in check_blocks):
        _vprint(f"插入目录 (深度: {opts.toc_depth})…")
        _add_toc(doc, ctx.styles, depth=opts.toc_depth)
        _toc_pending = False
        _toc_inserted = True

    if opts.number_headings:
        _vprint("启用标题自动编号…")
        ctx.reset_heading_counters()

    if fn_list:
        _vprint(f"添加 {len(fn_list)} 个脚注…")
        ctx.fn_map = add_footnotes_to_document(doc, fn_list)

    _vprint("处理内容块…")
    table_counter = [0]
    block_count = len(blocks)

    for idx, block in enumerate(blocks):
        tag = block.tag
        try:
            if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
                handle_heading(doc, block, int(tag[1]), ctx.styles, ctx,
                               number_headings=opts.number_headings,
                               page_break=opts.page_break_h1)
                if _toc_pending and tag == "h1":
                    _add_toc(doc, ctx.styles, depth=opts.toc_depth)
                    _toc_pending = False
                    _toc_inserted = True
            elif tag == "p":
                handle_paragraph(doc, block, ctx.styles, ctx,
                                 image_max_width=opts.image_max_width)
            elif tag == "img":
                handle_image(doc, block, ctx.styles, ctx,
                             image_max_width=opts.image_max_width)
                if opts.formula_numbering and block.get("class", "").startswith("math-block"):
                    from .ooxml_helpers import add_seq_number as _add_seq_number
                    p_num = doc.add_paragraph()
                    p_num.paragraph_format.alignment = 3
                    _add_seq_number(p_num, "Equation")
            elif tag == "blockquote":
                handle_blockquote(doc, block, ctx.styles, ctx)
            elif tag == "pre":
                handle_code_block(doc, block, ctx.styles, ctx,
                                  highlight_enabled=opts.highlight_enabled,
                                  mermaid_enabled=opts.mermaid_enabled)
            elif tag == "ul":
                handle_unordered_list(doc, block, ctx.styles, ctx)
            elif tag == "ol":
                handle_ordered_list(doc, block, ctx.styles, ctx)
            elif tag == "table":
                handle_table(doc, block, ctx.styles, ctx,
                             three_line=opts.three_line_table,
                             table_idx=table_counter[0])
                table_counter[0] += 1
            elif tag == "hr":
                handle_horizontal_rule(doc, ctx.styles)
            else:
                handle_paragraph(doc, block, ctx.styles, ctx,
                                 image_max_width=opts.image_max_width)
            _vprint(f"  [{idx+1}/{block_count}] {tag} ✓")
        except Exception as e:
            context_preview = _inline_text_preview(block, 80)
            ctx.report.warn(f"处理 <{tag}> 块时出错 ({context_preview}): {e}")

    if _toc_inserted:
        _vprint("提示：目录需右键 → 更新域后生效")

    _vprint("保存文档…")
    doc.save(str(output_path))

    _vprint("修复 OOXML 元数据…")
    fix_ooxml_metadata(output_path)

    ctx.report.info_msg("输出文件: " + str(output_path))
    print(f"  {ctx.report.summary()}")
    return ctx.report


def _inline_text_preview(elem, max_len: int = 80) -> str:
    """Return first *max_len* characters of inline text from *elem*."""
    from .handlers import inline_text
    text = inline_text(elem).strip()
    if len(text) > max_len:
        return text[:max_len] + "…"
    return text
