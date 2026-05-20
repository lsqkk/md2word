"""Template parsing — extract formatting from guide paragraphs in a .docx."""

from __future__ import annotations

import dataclasses
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Pt, RGBColor


def set_run_font(run, western: str, eastern: str | None = None) -> None:
    """Set both Western and East-Asian fonts on a run.

    This is critical for Chinese text — without an explicit *eastern* font
    Word may fall back to Japanese fonts (MS Mincho, etc.) via the theme.
    """
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), western)
    rFonts.set(qn("w:hAnsi"), western)
    target_ea = eastern or western
    rFonts.set(qn("w:eastAsia"), target_ea)


def _get_run_east_asia(run) -> str | None:
    """Return the East-Asian font name of a run, if set."""
    rPr = run._r.find(qn("w:rPr"))
    if rPr is None:
        return None
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        return None
    ea = rFonts.get(qn("w:eastAsia"))
    return ea


@dataclasses.dataclass
class ParagraphFormat:
    """Captures the formatting of a guide paragraph."""

    font_name: str | None = None
    east_asia_font: str | None = None
    font_size_emu: int | None = None
    bold: bool | None = None
    italic: bool | None = None
    color_rgb: tuple[int, int, int] | None = None
    alignment: int | None = None
    first_line_indent_emu: int | None = None
    left_indent_emu: int | None = None
    space_before_emu: int | None = None
    space_after_emu: int | None = None
    line_spacing: float | None = None

    @classmethod
    def from_docx_paragraph(cls, p) -> "ParagraphFormat":
        pf = cls()
        run = p.runs[0] if p.runs else None

        if run:
            if run.font.name:
                pf.font_name = run.font.name
            pf.east_asia_font = _get_run_east_asia(run)
            if run.font.size:
                pf.font_size_emu = run.font.size
            if run.font.bold is not None:
                pf.bold = run.font.bold
            if run.font.italic is not None:
                pf.italic = run.font.italic
            if run.font.color and run.font.color.rgb:
                pf.color_rgb = (
                    run.font.color.rgb[0],
                    run.font.color.rgb[1],
                    run.font.color.rgb[2],
                )

        ppf = p.paragraph_format
        if ppf.alignment is not None:
            pf.alignment = ppf.alignment
        if ppf.first_line_indent:
            pf.first_line_indent_emu = ppf.first_line_indent
        if ppf.left_indent:
            pf.left_indent_emu = ppf.left_indent
        if ppf.space_before:
            pf.space_before_emu = ppf.space_before
        if ppf.space_after:
            pf.space_after_emu = ppf.space_after
        if ppf.line_spacing:
            pf.line_spacing = ppf.line_spacing

        return pf

    def apply_to_paragraph(self, p) -> None:
        ppf = p.paragraph_format
        if self.alignment is not None:
            ppf.alignment = self.alignment
        if self.first_line_indent_emu is not None:
            ppf.first_line_indent = Emu(self.first_line_indent_emu)
        if self.left_indent_emu is not None:
            ppf.left_indent = Emu(self.left_indent_emu)
        if self.space_before_emu is not None:
            ppf.space_before = Emu(self.space_before_emu)
        if self.space_after_emu is not None:
            ppf.space_after = Emu(self.space_after_emu)
        if self.line_spacing is not None:
            ppf.line_spacing = self.line_spacing

    def apply_to_run(self, run) -> None:
        w = self.font_name or "SimSun"
        ea = self.east_asia_font or w
        set_run_font(run, w, ea)
        if self.font_size_emu is not None:
            run.font.size = Emu(self.font_size_emu)
        if self.bold is not None:
            run.font.bold = self.bold
        if self.italic is not None:
            run.font.italic = self.italic
        if self.color_rgb is not None:
            run.font.color.rgb = RGBColor(*self.color_rgb)


_STYLE_KEYWORDS: dict[str, str] = {
    "一级标题": "h1",
    "heading 1": "h1",
    "二级标题": "h2",
    "heading 2": "h2",
    "三级标题": "h3",
    "heading 3": "h3",
    "四级标题": "h4",
    "heading 4": "h4",
    "五级标题": "h5",
    "heading 5": "h5",
    "正文": "body",
    "首行缩进": "body_indent",
    "图片": "image",
    "图注": "figcaption",
    "引用": "quote",
    "代码": "code",
    "code block": "code",
    "无序列表": "bullet_list",
    "bullet list": "bullet_list",
    "有序列表": "number_list",
    "ordered list": "number_list",
    "目录标题": "toc_title",
    "table of contents": "toc_title",
}


def _guess_slot(text: str) -> str | None:
    """Match guide paragraph text to a style slot.

    Strips whitespace and punctuation then does substring matching,
    reducing false positives from incidental keyword inclusion.
    """
    import re
    normalised = re.sub(r"[^a-zA-Z0-9一-鿿]", "", text.lower().strip())
    for keyword in sorted(_STYLE_KEYWORDS, key=len, reverse=True):
        kw_norm = re.sub(r"[^a-zA-Z0-9一-鿿]", "", keyword.lower())
        if kw_norm and kw_norm in normalised:
            return _STYLE_KEYWORDS[keyword]
    return None


def extract_template_styles(template_path: str | Path) -> dict[str, ParagraphFormat]:
    doc = Document(str(template_path))
    styles: dict[str, ParagraphFormat] = {}

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        slot = _guess_slot(text)
        if slot and slot not in styles:
            styles[slot] = ParagraphFormat.from_docx_paragraph(p)

    if "body" not in styles:
        styles["body"] = ParagraphFormat()
    return styles


def validate_template(template_path: str | Path) -> dict[str, list[str]]:
    """Validate a template and return missing/redundant guide slots."""
    doc = Document(str(template_path))
    found: set[str] = set()
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        slot = _guess_slot(text)
        if slot:
            found.add(slot)

    required = {"h1", "h2", "h3", "body", "code"}
    recommended = {"body_indent", "image", "figcaption", "quote", "bullet_list", "number_list"}

    return {
        "missing_required": sorted(required - found),
        "missing_recommended": sorted(recommended - found),
        "found": sorted(found),
    }


def list_template_styles(template_path: str | Path) -> list[dict[str, Any]]:
    doc = Document(str(template_path))
    items: list[dict[str, Any]] = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        slot = _guess_slot(text)
        if slot:
            fmt = ParagraphFormat.from_docx_paragraph(p)
            ea = fmt.east_asia_font or "(same)"
            items.append(
                {
                    "slot": slot,
                    "guide_text": text,
                    "font": fmt.font_name or "(inherited)",
                    "ea_font": ea,
                    "size_pt": (
                        round(fmt.font_size_emu / 12700, 1)
                        if fmt.font_size_emu
                        else "(inherited)"
                    ),
                    "bold": fmt.bold,
                    "align": str(fmt.alignment),
                }
            )
    return items
